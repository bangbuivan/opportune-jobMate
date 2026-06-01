# -*- coding: utf-8 -*-
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import sys
import os

def set_cell_shading(cell, color_hex):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_margins(cell, top=120, bottom=120, left=160, right=160):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)
    
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4') 
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'CCCCCC') 
        tblBorders.append(border)

def set_cell_text(cell, text, bold=False, italic=False, size_pt=10.5, color_rgb=RGBColor(0,0,0), align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size_pt)
    run.font.color.rgb = color_rgb
    run.bold = bold
    run.italic = italic

def style_table(table, header_color="1A5276", row_colors=["F8F9F9", "FFFFFF"]):
    for cell in table.rows[0].cells:
        set_cell_shading(cell, header_color)
        set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11)

    for r_idx, row in enumerate(table.rows[1:], start=1):
        bg_color = row_colors[(r_idx - 1) % len(row_colors)]
        for cell in row.cells:
            if bg_color:
                set_cell_shading(cell, bg_color)
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(10.5)

def build_report():
    doc = Document()
    
    # 1. Page Configuration (Standard A4, 1-inch margins)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.page_width = Inches(8.27)  
        section.page_height = Inches(11.69) 

    # 2. Modify Default Style (Normal)
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Times New Roman'
    normal_font.size = Pt(12)
    normal_font.color.rgb = RGBColor(0, 0, 0)
    normal_style.paragraph_format.line_spacing = 1.3
    normal_style.paragraph_format.space_after = Pt(6)
    normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 3. Configure Heading Styles
    def config_heading(style_name, size, color, bold=True, italic=False, before=12, after=6):
        style = doc.styles[style_name]
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(size)
        font.bold = bold
        font.italic = italic
        font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    config_heading('Heading 1', 16, RGBColor(26, 82, 118), before=18, after=8)   
    config_heading('Heading 2', 13, RGBColor(33, 97, 140), before=14, after=6)   
    config_heading('Heading 3', 12, RGBColor(44, 62, 80), italic=True, before=10, after=4) 

    # --- COVER PAGE ---
    p_top = doc.add_paragraph()
    p_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_top1 = p_top.add_run("ĐẠI HỌC PHENIKAA\nTRƯỜNG CÔNG NGHỆ THÔNG TIN PHENIKAA\n")
    run_top1.font.size = Pt(13)
    run_top1.bold = True
    run_top2 = p_top.add_run("-------------------\n\n\n\n")
    run_top2.font.size = Pt(11)
    
    p_type = doc.add_paragraph()
    p_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_type.paragraph_format.space_after = Pt(18)
    run_type = p_type.add_run("BÁO CÁO ĐỒ ÁN CƠ SỞ")
    run_type.font.size = Pt(15)
    run_type.bold = True
    run_type.font.color.rgb = RGBColor(120, 120, 120)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(110)
    run_title = p_title.add_run(
        "TÊN ĐỀ TÀI: PHÁT TRIỂN HỆ THỐNG HỖ TRỢ HỒ SƠ NGHỀ NGHIỆP THÔNG MINH\n"
        "VÀ MÔ PHỎNG PHỎNG VẤN TÍCH HỢP TRÍ TUỆ NHÂN TẠO\n"
        "(CAREERFORGE AI)"
    )
    run_title.font.size = Pt(18)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor(26, 82, 118)

    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_info.paragraph_format.left_indent = Inches(1.5)
    p_info.paragraph_format.space_after = Pt(90)
    
    r_info = p_info.add_run(
        "Học phần:\t\tĐồ Án Cơ Sở CSE702013-1-2-25(N01)\n"
        "Giảng viên hướng dẫn:\tThs. Vũ Văn Quang\n"
        "Sinh viên thực hiện:\tBùi Văn Bằng  - MSV: 22014000\n"
        "\t\t\tTạ Văn Thanh - MSV: 22010161\n"
        "Ngành học:\t\tCông nghệ Thông tin"
    )
    r_info.font.size = Pt(13)
    
    p_bot = doc.add_paragraph()
    p_bot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_bot = p_bot.add_run("HÀ NỘI - 2026")
    r_bot.font.size = Pt(12)
    r_bot.bold = True

    doc.add_page_break()

    # --- LỜI MỞ ĐẦU ---
    p_preface_h = doc.add_paragraph()
    p_preface_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_preface_h.paragraph_format.space_before = Pt(12)
    p_preface_h.paragraph_format.space_after = Pt(12)
    run_preface_h = p_preface_h.add_run("LỜI MỞ ĐẦU")
    run_preface_h.font.size = Pt(14)
    run_preface_h.bold = True
    run_preface_h.font.color.rgb = RGBColor(26, 82, 118)

    p_preface_1 = doc.add_paragraph(
        "Trong bối cảnh công nghệ thông tin và internet ngày càng phát triển mạnh mẽ, việc ứng dụng trí tuệ nhân tạo (AI) và xử lý ngôn ngữ tự nhiên (NLP) "
        "vào quy trình tuyển dụng đã trở thành một xu thế tất yếu của các doanh nghiệp trên toàn thế giới. Hầu hết các nhà tuyển dụng công nghệ lớn hiện nay "
        "đều sử dụng các Hệ thống quản lý và sàng lọc hồ sơ tự động (Applicant Tracking System - ATS) để quét từ khóa, lọc năng lực và đánh giá sơ bộ CV ứng viên. "
        "Điều này đặt ra một thách thức lớn đối với sinh viên mới tốt nghiệp và người tìm việc ngành IT: Làm sao để viết một bản CV đúng cấu trúc, tối ưu hóa từ khóa "
        "kỹ thuật và định lượng hóa thành tựu để vượt qua bộ lọc ATS trước khi tiếp cận được với nhà tuyển dụng."
    )
    p_preface_2 = doc.add_paragraph(
        "Nhận thức được xu thế đó, cùng với những kiến thức đã tích lũy trong quá trình học tập tại Trường Công nghệ Thông tin Phenikaa, nhóm chúng em quyết định chọn "
        "đề tài \"Phát triển hệ thống hỗ trợ hồ sơ nghề nghiệp thông minh và mô phỏng phỏng vấn tích hợp trí tuệ nhân tạo (CareerForge AI)\" cho đồ án cơ sở của mình. "
        "Hệ thống hướng tới việc cung cấp một bộ công cụ toàn diện giúp người dùng phân tích lỗi kỹ thuật CV, tính toán tương hợp CV-JD, gợi ý lộ trình học tập để bù đắp "
        "khoảng trống kỹ năng, đồng thời tạo một phòng giả lập phỏng vấn tương tác thông minh dựa trên mô hình ngôn ngữ lớn (LLM)."
    )
    p_preface_3 = doc.add_paragraph(
        "Để hoàn thành được đồ án này, chúng em xin gửi lời cảm ơn chân thành và sâu sắc nhất đến Ths. Vũ Văn Quang. Sự hướng dẫn tận tình, những định hướng học thuật "
        "quý báu và những ý kiến phản biện sắc sảo của thầy trong suốt học phần Đồ Án Cơ Sở đã giúp chúng em hoàn thiện hệ thống và giải quyết được các bài toán kỹ thuật phức tạp."
    )
    p_preface_4 = doc.add_paragraph(
        "Mặc dù đã nỗ lực hết sức, song do kiến thức, kinh nghiệm thực tế còn hạn chế và thời gian triển khai có hạn, đồ án chắc chắn không thể tránh khỏi những thiếu sót "
        "về mặt tối ưu hóa kiến trúc hoặc trải nghiệm người dùng. Chúng em rất mong nhận được những ý kiến đóng góp quý báu từ thầy/cô để đề tài này được phát triển hoàn thiện hơn."
    )

    doc.add_page_break()

    # --- BẢNG PHÂN CHIA CÔNG VIỆC ---
    p_work_h = doc.add_paragraph()
    p_work_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_work_h.paragraph_format.space_before = Pt(12)
    p_work_h.paragraph_format.space_after = Pt(12)
    run_work_h = p_work_h.add_run("BẢNG PHÂN CHIA CÔNG VIỆC")
    run_work_h.font.size = Pt(14)
    run_work_h.bold = True
    run_work_h.font.color.rgb = RGBColor(26, 82, 118)

    table_work = doc.add_table(rows=3, cols=3)
    table_work.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table_work)
    
    headers_work = ["Nhiệm vụ chính", "Thành viên thực hiện", "Đóng góp (%)"]
    for i, h_text in enumerate(headers_work):
        set_cell_text(table_work.rows[0].cells[i], h_text, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
        
    set_cell_text(table_work.rows[1].cells[0], "Kiến trúc hệ thống, Lập trình FastAPI Backend, Xử lý spaCy NLP & RapidFuzz Engine, Gợi ý FAISS, Tích hợp Gemini API, Viết tài liệu.")
    set_cell_text(table_work.rows[1].cells[1], "Bùi Văn Bằng\n(MSV: 22014000)")
    set_cell_text(table_work.rows[1].cells[2], "50%")

    set_cell_text(table_work.rows[2].cells[0], "Thiết kế UI/UX Streamlit Frontend, Quản lý Session State, Xây dựng JobRadar, JobMatcher, Cover Letter & Mock Interview, Lập trình Word Generator, Kiểm thử hệ thống.")
    set_cell_text(table_work.rows[2].cells[1], "Tạ Văn Thanh\n(MSV: 22010161)")
    set_cell_text(table_work.rows[2].cells[2], "50%")
    
    style_table(table_work)

    doc.add_page_break()

    # --- MỤC LỤC ---
    doc.add_heading("MỤC LỤC", level=1)
    p_toc = doc.add_paragraph()
    p_toc.paragraph_format.line_spacing = 1.4
    
    toc_entries = [
        ("CHƯƠNG 1: GIỚI THIỆU", "1"),
        ("Đặt vấn đề", "1"),
        ("Xác định nhiệm vụ, mục đích và phạm vi của đề tài", "1"),
        ("CHƯƠNG 2: PHÂN TÍCH THIẾT KẾ HỆ THỐNG", "3"),
        ("2.1. Nội dung công việc", "3"),
        ("2.2. Phân tích thiết kế", "4"),
        ("2.3. Phân tích chức năng hệ thống", "5"),
        ("2.3.1. Chức năng của guest", "5"),
        ("2.3.2. Chức năng của user", "6"),
        ("2.3.3. Chức năng của admin", "6"),
        ("2.4. Biểu đồ Usecase và biểu đồ hoạt động", "7"),
        ("2.4.1. Biểu đồ Usecase tổng quát", "7"),
        ("2.4.2. Biểu đồ hoạt động phân tích ATS (ATS Tune-Up)", "7"),
        ("2.4.3. Biểu đồ hoạt động so khớp CV-JD (JobMatcher)", "8"),
        ("2.4.4. Biểu đồ hoạt động gợi ý việc làm (CareerMatch)", "8"),
        ("2.4.5. Biểu đồ hoạt động cầu nối kỹ năng (SkillBridge)", "9"),
        ("2.4.6. Biểu đồ hoạt động tạo và xuất bản CV (ResumeBuilder)", "9"),
        ("2.4.7. Biểu đồ hoạt động luyện tập phỏng vấn thử (Mock Interview)", "10"),
        ("2.5. Biểu đồ tuần tự", "10"),
        ("2.5.1. Biểu đồ tuần tự chức năng phân tích ATS", "10"),
        ("2.5.2. Biểu đồ tuần tự chức năng so khớp và gợi ý việc làm", "11"),
        ("2.5.3. Biểu đồ tuần tự chức năng luyện phỏng vấn thử", "11"),
        ("2.6. Tổng quan về cơ sở dữ liệu", "12"),
        ("CHƯƠNG 3: CÔNG NGHỆ SỬ DỤNG", "14"),
        ("3.1. Công nghệ phía frontend", "14"),
        ("3.1.1. Nền tảng và điều hướng", "14"),
        ("3.1.2. Quản lý trạng thái và dữ liệu", "15"),
        ("3.1.3. Thiết kế giao diện", "15"),
        ("3.1.4. Các tiện ích bổ sung", "15"),
        ("3.2. Công nghệ phía backend và database", "16"),
        ("3.2.1. Nền tảng và framework", "16"),
        ("3.2.2. Database", "18"),
        ("3.2.3. Bảo mật và phân quyền", "18"),
        ("3.3. Mã nguồn", "19"),
        ("CHƯƠNG 4: TRIỂN KHAI VÀ THỰC NGHIỆM", "20"),
        ("4.1. Giao diện trang chủ và Dashboard Skill Market", "20"),
        ("4.2. Giao diện tìm kiếm và định hướng việc làm (JobRadar)", "20"),
        ("4.3. Giao diện so khớp CV với mô tả công việc (JobMatcher)", "21"),
        ("4.4. Giao diện kết quả phân tích kỹ năng trùng và thiếu từ JobMatcher", "21"),
        ("4.5. Giao diện gợi ý nghề nghiệp tự động (CareerMatch)", "22"),
        ("4.6. Giao diện cầu nối kỹ năng và biểu đồ mạng nhện (SkillBridge)", "22"),
        ("4.7. Giao diện lộ trình tự học đề xuất bởi AI từ SkillBridge", "23"),
        ("4.8. Giao diện biểu mẫu động tạo CV (ResumeBuilder)", "23"),
        ("4.9. Giao diện tối ưu hóa mô tả công việc theo cấu trúc STAR bằng AI", "24"),
        ("4.10. Giao diện tải xuống tệp CV Word (.docx) chuẩn ATS", "24"),
        ("4.11. Giao diện đánh giá ATS tĩnh (ATS Tune-Up Local)", "25"),
        ("4.12. Giao diện phân tích chất lượng CV bằng AI (ATS Tune-Up AI)", "25"),
        ("4.13. Giao diện soạn thảo Cover Letter cá nhân hóa bằng AI", "26"),
        ("4.14. Giao diện tải xuống tệp Cover Letter (.docx) trang trọng", "26"),
        ("4.15. Giao diện giả lập phòng phỏng vấn thử (Mock Interview)", "27"),
        ("4.16. Giao diện kết quả chấm điểm phỏng vấn tổng quan từ AI", "27"),
        ("4.17. Giao diện phân tích câu trả lời chi tiết và câu trả lời mẫu từ AI", "28"),
        ("4.18. Giao diện tài liệu API tự động (FastAPI Swagger UI)", "28"),
        ("4.19. Những hạn chế và hướng phát triển của dự án", "29"),
        ("TÀI LIỆU THAM KHẢO", "30")
    ]
    
    for title, pg in toc_entries:
        p_item = doc.add_paragraph()
        p_item.paragraph_format.space_after = Pt(2)
        if title.startswith("2.") or title.startswith("3.") or title.startswith("4.") or title.startswith("1.") or title.startswith("Đặt") or title.startswith("Xác"):
            if title.startswith("2.3.") or title.startswith("2.4.") or title.startswith("2.5.") or title.startswith("3.1.") or title.startswith("3.2."):
                p_item.paragraph_format.left_indent = Inches(0.6)
            else:
                p_item.paragraph_format.left_indent = Inches(0.3)
            run = p_item.add_run(title)
        else:
            run = p_item.add_run(title)
            run.bold = True
        
        dots_count = 90 - len(title)
        if dots_count < 5: dots_count = 5
        p_item.add_run(" ." * dots_count)
        
        run_pg = p_item.add_run(f" {pg}")
        run_pg.bold = True

    doc.add_page_break()

    # --- DANH MỤC HÌNH ẢNH ---
    doc.add_heading("DANH MỤC HÌNH ẢNH", level=1)
    p_fig = doc.add_paragraph()
    p_fig.paragraph_format.line_spacing = 1.3
    
    fig_entries = [
        ("Hình 1. Các tác nhân và kiến trúc tổng quan hệ thống", "5"),
        ("Hình 2. Biểu đồ Usecase tổng quát hệ thống CareerForge AI", "7"),
        ("Hình 3. Biểu đồ hoạt động trích xuất thông tin CV (CV Parsing)", "7"),
        ("Hình 4. Biểu đồ hoạt động phân tích ATS (ATS Tune-Up)", "8"),
        ("Hình 5. Biểu đồ hoạt động so khớp CV-JD (JobMatcher)", "8"),
        ("Hình 6. Biểu đồ hoạt động gợi ý việc làm (CareerMatch)", "9"),
        ("Hình 7. Biểu đồ hoạt động tạo và xuất bản CV (ResumeBuilder)", "9"),
        ("Hình 8. Biểu đồ hoạt động luyện tập phỏng vấn thử (Mock Interview)", "10"),
        ("Hình 9. Biểu đồ tuần tự chức năng phân tích ATS (ATS Tune-Up)", "10"),
        ("Hình 10. Biểu đồ tuần tự chức năng so khớp và gợi ý việc làm", "11"),
        ("Hình 11. Biểu đồ tuần tự chức năng luyện phỏng vấn thử", "11"),
        ("Hình 12. Giao diện trang chủ và Dashboard Skill Market", "20"),
        ("Hình 13. Giao diện tìm kiếm và định hướng việc làm (JobRadar)", "20"),
        ("Hình 14. Giao diện so khớp CV với mô tả công việc (JobMatcher)", "21"),
        ("Hình 15. Giao diện kết quả phân tích kỹ năng trùng và thiếu từ JobMatcher", "21"),
        ("Hình 16. Giao diện gợi ý nghề nghiệp tự động (CareerMatch)", "22"),
        ("Hình 17. Giao diện cầu nối kỹ năng và biểu đồ mạng nhện (SkillBridge)", "22"),
        ("Hình 18. Giao diện lộ trình tự học đề xuất bởi AI từ SkillBridge", "23"),
        ("Hình 19. Giao diện biểu mẫu động tạo CV (ResumeBuilder)", "23"),
        ("Hình 20. Giao diện tối ưu hóa mô tả công việc theo cấu trúc STAR bằng AI", "24"),
        ("Hình 21. Giao diện tải xuống tệp CV Word (.docx) chuẩn ATS", "24"),
        ("Hình 22. Giao diện đánh giá ATS tĩnh (ATS Tune-Up Local)", "25"),
        ("Hình 23. Giao diện phân tích chất lượng CV bằng AI (ATS Tune-Up AI)", "25"),
        ("Hình 24. Giao diện soạn thảo Cover Letter cá nhân hóa bằng AI", "26"),
        ("Hình 25. Giao diện tải xuống tệp Cover Letter (.docx) trang trọng", "26"),
        ("Hình 26. Giao diện giả lập phòng phỏng vấn thử (Mock Interview)", "27"),
        ("Hình 27. Giao diện kết quả chấm điểm phỏng vấn tổng quan từ AI", "27"),
        ("Hình 28. Giao diện phân tích câu trả lời chi tiết và câu trả lời mẫu từ AI", "28"),
        ("Hình 29. Giao diện tài liệu API tự động (FastAPI Swagger UI)", "28"),
    ]
    
    for title, pg in fig_entries:
        p_item = doc.add_paragraph()
        p_item.paragraph_format.space_after = Pt(2)
        p_item.paragraph_format.left_indent = Inches(0.3)
        run = p_item.add_run(title)
        
        dots_count = 90 - len(title)
        if dots_count < 5: dots_count = 5
        p_item.add_run(" ." * dots_count)
        
        run_pg = p_item.add_run(f" {pg}")
        run_pg.bold = True

    doc.add_page_break()

    # --- CHƯƠNG 1 ---
    doc.add_heading("CHƯƠNG 1: GIỚI THIỆU", level=1)
    
    doc.add_heading("Đặt vấn đề:", level=2)
    p = doc.add_paragraph(
        "Trong kỷ nguyên số, ngành Công nghệ thông tin (IT) đã trở thành động cơ cốt lõi thúc đẩy sự phát triển của nền kinh tế toàn cầu. "
        "Tại Việt Nam, sự phát triển nhanh chóng của các doanh nghiệp phần mềm và làn sóng chuyển dịch công nghệ tạo ra nhu cầu tuyển dụng lập trình viên cực kỳ lớn. "
        "Tuy nhiên, thị trường lao động IT cũng chứng kiến sự sàng lọc tự nhiên vô cùng gay gắt. Đối với một người tìm việc, thách thức lớn nhất "
        "không chỉ là tích lũy kiến thức chuyên môn mà còn nằm ở cách trình bày năng lực đó hiệu quả trên hồ sơ cá nhân (CV). "
        "Hầu hết các doanh nghiệp công nghệ lớn hiện nay đều áp dụng Hệ thống quản lý và sàng lọc hồ sơ tự động (Applicant Tracking System - ATS) "
        "như một công cụ gác cổng đầu tiên nhằm giảm thiểu chi phí và thời gian cho phòng tuyển dụng."
    )
    p = doc.add_paragraph(
        "ATS hoạt động dựa trên các thuật toán phân tích văn bản để bóc tách các trường dữ liệu và chấm điểm mức độ phù hợp của CV dựa trên từ khóa kỹ năng "
        "so với bản mô tả công việc (JD). Thống kê tuyển dụng chỉ ra rằng hơn 75% CV bị loại bỏ ngay lập tức ở vòng này vì các lỗi định dạng kỹ thuật "
        "như bố cục chia cột phức tạp, chèn bảng biểu hoặc hình ảnh gây lỗi trích xuất văn bản thô, thiếu các từ khóa công nghệ bắt buộc, lạm dụng đại từ nhân xưng "
        "hoặc viết câu thiếu định lượng thành tựu. Ứng viên dù có chuyên môn tốt nhưng thiếu kiến thức tối ưu hóa CV theo chuẩn ATS vẫn mất đi cơ hội tiếp cận nhà tuyển dụng."
    )
    p = doc.add_paragraph(
        "Đồng thời, sau khi vượt qua vòng hồ sơ, ứng viên thường thiếu môi trường thực hành phỏng vấn thực tế. Việc chuẩn bị phỏng vấn thường bị thụ động "
        "do ứng viên chỉ đọc các câu hỏi lý thuyết chung chung trên mạng mà không được luyện tập trả lời tình huống cá nhân hóa sát với hồ sơ của mình "
        "và JD ứng tuyển. Chính vì những lý do đó, dự án “CareerForge AI” được nghiên cứu và phát triển để cung cấp một hệ sinh thái hỗ trợ hồ sơ nghề nghiệp "
        "thông minh, tối ưu hóa CV chuẩn ATS và mô phỏng phỏng vấn tương tác bằng trí tuệ nhân tạo, hỗ trợ tối đa cho sinh viên và lập trình viên IT."
    )

    doc.add_heading("Xác định nhiệm vụ, mục đích và phạm vi của đề tài:", level=2)
    p = doc.add_paragraph()
    p.add_run("Nhiệm vụ:\n").bold = True
    p.add_run(
        "+) Nghiên cứu và xây dựng bộ trích xuất thông tin CV và JD sử dụng mô hình spaCy NLP kết hợp giải thuật so khớp từ khóa mờ RapidFuzz.\n"
        "+) Thiết kế và phát triển thuật toán tìm kiếm vector ngữ nghĩa sử dụng Sentence Transformers và cơ sở dữ liệu vector FAISS, kết hợp tìm kiếm từ khóa TF-IDF để gợi ý công việc tương thích.\n"
        "+) Tích hợp Google Gemini API để thực thi các tính năng thông minh: phân tích lỗi CV nâng cao, sinh câu hỏi phỏng vấn cá nhân hóa và chấm điểm tương tác.\n"
        "+) Lập trình module kết xuất tệp Word (.docx) chuyên nghiệp đáp ứng các tiêu chuẩn ATS qua 4 phong cách giao diện."
    )
    
    p = doc.add_paragraph()
    p.add_run("Mục đích:\n").bold = True
    p.add_run(
        "+) Đáp ứng nhu cầu chuẩn bị hồ sơ ngày càng phát triển của sinh viên ngành công nghệ, xây dựng lên một công cụ hỗ trợ nghề nghiệp hiệu quả.\n"
        "+) Giúp người lao động ngành IT tự đánh giá chất lượng hồ sơ, định vị khoảng trống kỹ năng để lập lộ trình học tập nâng cao năng lực.\n"
        "+) Hỗ trợ ứng viên viết CV chuẩn hóa, tối ưu hóa từ khóa chuyên môn để tăng tỷ lệ phản hồi từ bộ lọc tuyển dụng.\n"
        "+) Tạo môi trường luyện tập phỏng vấn thử có phản hồi sửa lỗi tức thì giúp ứng viên tự tin trước buổi gặp mặt chính thức."
    )

    p = doc.add_paragraph()
    p.add_run("Phạm vi:\n").bold = True
    p.add_run(
        "+) Đề tài sẽ tập trung hỗ trợ các vai trò phổ biến trong ngành công nghiệp phần mềm: Frontend, Backend, Fullstack, Mobile Developer, DevOps Engineer, và Data Specialist.\n"
        "+) Đối tượng sử dụng chính là các lập trình viên, sinh viên ngành CNTT có nhu cầu tìm việc hoặc tối ưu hóa hồ sơ năng lực cá nhân.\n"
        "+) Ứng dụng hoạt động dưới dạng Web App chạy cục bộ hoặc triển khai trên cloud, xử lý bảo mật dữ liệu hoàn toàn trong bộ nhớ đệm (In-memory buffer)."
    )

    doc.add_page_break()

    # --- CHƯƠNG 2 ---
    doc.add_heading("CHƯƠNG 2: PHÂN TÍCH THIẾT KẾ HỆ THỐNG", level=1)
    
    doc.add_heading("2.1. Nội dung công việc:", level=2)
    p = doc.add_paragraph(
        "Hệ thống CareerForge AI tập trung xử lý dữ liệu hồ sơ CV của ứng viên và bản mô tả công việc (JD) của nhà tuyển dụng. Nội dung công việc bao gồm "
        "tiền xử lý văn bản, trích xuất thực thể kỹ năng, tính toán độ tương hợp ngữ nghĩa và tích hợp mô hình ngôn ngữ lớn (LLM). "
        "Công việc được phân tách thành 3 mảng nghiệp vụ chính: 1) Tối ưu hóa hồ sơ tĩnh qua 12 tiêu chí ATS; 2) Hệ thống gợi ý và so khớp "
        "việc làm (Hybrid Search); 3) Giả lập và đánh giá phỏng vấn tương tác thông minh."
    )
    p = doc.add_paragraph(
        "Trong đó, trọng tâm cốt lõi của việc tối ưu hóa hồ sơ tĩnh là quét và đánh giá CV dựa trên 12 tiêu chí kỹ thuật ATS được định nghĩa chi tiết như sau:"
    )
    
    # Detail the 12 ATS criteria
    ats_items = [
        ("1. Kiểm tra Thông tin Liên hệ (Contact Information)", "Phát hiện sự hiện diện của Email, Số điện thoại và liên kết hồ sơ LinkedIn của ứng viên bằng cách áp dụng các biểu thức chính quy (Regular Expressions). Đây là tiêu chí tiên quyết vì các hệ thống ATS sẽ tự động loại bỏ các CV thiếu thông tin liên lạc."),
        ("2. Kiểm tra Chính tả & Ngữ pháp (Spelling & Grammar)", "Kiểm tra lỗi viết hoa hoặc lỗi chính tả cơ bản dựa trên cấu trúc các thực thể và chữ viết hoa không nhất quán. Đảm bảo văn bản CV chuyên nghiệp, không có các lỗi soạn thảo ngớ ngẩn."),
        ("3. Kiểm tra Đại từ Nhân xưng (Personal Pronouns)", "Tìm kiếm các đại từ như 'I', 'me', 'my', 'we', 'our'. Theo chuẩn viết CV chuyên nghiệp của ATS, ứng viên phải tránh sử dụng đại từ nhân xưng để duy trì tính khách quan và tập trung vào hành động dạng liệt kê."),
        ("4. Nhắm mục tiêu Kỹ năng & Từ khóa (Skills Targeting)", "Trích xuất kỹ năng cứng và kỹ năng mềm bằng spaCy NLP và đối chiếu với từ điển kỹ năng. ATS chấm điểm mật độ từ khóa rất cao, nên việc phát hiện kỹ năng tương ứng với vị trí là bắt buộc."),
        ("5. Kiểm tra Câu quá dài (Complex or Long Sentences)", "Nhận diện các câu văn có độ dài vượt quá 40 từ thông qua việc phân tích cấu trúc cây cú pháp của spaCy. Các câu quá dài thường gây khó khăn cho bộ parser của ATS và làm giảm trải nghiệm của nhà tuyển dụng."),
        ("6. Kiểm tra Cụm từ chung chung (Generic Phrases)", "Đếm tần suất xuất hiện của các cụm từ sáo rỗng, yếu hoặc chung chung như 'responsible for', 'worked on', 'helped', 'involved in'. Nếu tần suất xuất hiện quá cao (>2 lần), hệ thống sẽ cảnh báo yêu cầu đổi sang động từ hành động mạnh mẽ hơn."),
        ("7. Nhận diện Câu bị động (Passive Voice)", "Xác định các cấu trúc bị động bằng cách quét liên động từ kết hợp với phân từ quá khứ (ví dụ: 'is developed by', 'was managed'). Câu bị động làm giảm tính chủ động và thuyết phục của hồ sơ, cần chuyển sang câu chủ động."),
        ("8. Điểm số định lượng (Quantified Achievements)", "Quét sự xuất hiện của các số liệu phần trăm (%), con số tài chính hoặc kết quả đo lường cụ thể (như 'tăng 30% hiệu năng', 'tiết kiệm 5000 USD'). ATS đánh giá rất cao các thành tựu được lượng hóa."),
        ("9. Các phần thiết yếu của CV (Essential Resume Sections)", "Xác nhận CV có đầy đủ 4 phần cốt lõi: 'Summary' (Tóm tắt), 'Education' (Học vấn), 'Experience' (Kinh nghiệm), 'Skills' (Kỹ năng). Sự thiếu sót bất kỳ phần nào sẽ khiến ATS phân loại sai thông tin."),
        ("10. Động từ hành động lặp lại (Repeated Action Verbs)", "Đếm số lần lặp lại của các động từ hành động chính (như 'developed', 'managed', 'led'). Việc lặp lại một từ quá nhiều lần thể hiện vốn từ vựng nghèo nàn, cần đa dạng hóa bằng từ đồng nghĩa."),
        ("11. Thuộc tính tài liệu (Document Properties)", "Kiểm tra số từ (yêu cầu trong khoảng 350-900 từ để tối ưu độ dài), định dạng tệp tin (khuyến nghị tệp PDF hoặc DOCX sạch) và kích thước tệp (dưới 2MB để hệ thống xử lý nhanh chóng)."),
        ("12. Sự nhất quán trong định dạng (Bullet Consistency)", "Kiểm tra việc sử dụng các ký tự đầu dòng (bullet points) như '•', '-', '●'. Định dạng bullet không nhất quán sẽ phá vỡ cấu trúc phân tích văn bản thô của ATS.")
    ]
    
    for title_ats, desc_ats in ats_items:
        p_ats = doc.add_paragraph()
        p_ats.paragraph_format.left_indent = Inches(0.3)
        p_ats.add_run(f"• {title_ats}: ").bold = True
        p_ats.add_run(desc_ats)

    doc.add_heading("2.2. Phân tích thiết kế:", level=2)
    p = doc.add_paragraph(
        "Hệ thống được thiết kế theo kiến trúc phân lớp hướng dịch vụ để đảm bảo tính module và khả năng bảo trì. Luồng dữ liệu tổng quát đi từ "
        "tệp tin người dùng tải lên thông qua bộ máy trích xuất văn bản thô (PyMuPDF/docx). Sau đó, chuỗi văn bản được đưa qua bộ NLP Engine "
        "để trích xuất kỹ năng chuyên môn. Dữ liệu này tiếp tục được sử dụng làm đầu vào cho các giải thuật tính điểm "
        "cục bộ (Rule-based), giải thuật tìm kiếm vector (FAISS + TF-IDF) và các dịch vụ đám mây (Google Gemini API) để trả về các phân tích chi tiết."
    )
    p = doc.add_paragraph(
        "Trong đó, động cơ NLP xử lý ngôn ngữ tự nhiên sử dụng mô hình spaCy kết hợp giải thuật so khớp từ khóa mờ RapidFuzz:"
    )
    
    p_nlp = doc.add_paragraph()
    p_nlp.paragraph_format.left_indent = Inches(0.3)
    p_nlp.add_run("• Động cơ xử lý ngôn ngữ tự nhiên spaCy:\n").bold = True
    p_nlp.add_run(
        "  Sử dụng mô hình ngôn ngữ pre-trained 'en_core_web_sm' để thực hiện phân tích cú pháp chuỗi văn bản CV và JD thô. "
        "Quy trình xử lý bao gồm:\n"
        "  1) Tách từ (Tokenization): Phân rã chuỗi văn bản thành các token đơn lẻ (từ, cụm từ, dấu câu).\n"
        "  2) Gán nhãn từ loại (Part-of-Speech Tagging - POS): Xác định vai trò ngữ pháp của từng từ (như Danh từ - NOUN, Động từ - VERB, Tính từ - ADJ). "
        "Thông tin này rất quan trọng để hệ thống phát hiện chính xác các động từ hành động và danh từ kỹ năng chuyên môn.\n"
        "  3) Nhận diện thực thể liên kết (Named Entity Recognition - NER): Phát hiện các thực thể như tên người, email, số điện thoại, tên trường đại học."
    )
    
    p_rf = doc.add_paragraph()
    p_rf.paragraph_format.left_indent = Inches(0.3)
    p_rf.add_run("• Giải thuật so khớp từ khóa mờ RapidFuzz:\n").bold = True
    p_rf.add_run(
        "  Sử dụng để giải quyết bài toán biến thể của từ khóa kỹ năng IT (như ứng viên viết 'ReactJS' nhưng JD yêu cầu 'React.js' hoặc 'React-JS'). "
        "Thuật toán dựa trên khoảng cách Levenshtein (Levenshtein Distance) đo số phép biến đổi ký tự tối thiểu để chuyển từ chuỗi này thành chuỗi kia. "
        "RapidFuzz áp dụng phương pháp 'Token Sort Ratio': Tách các chuỗi kỹ năng thành danh sách từ, sắp xếp theo thứ tự bảng chữ cái, và tính điểm tương đồng. "
        "Nhờ đó, hệ thống nhận diện chính xác các từ khóa công nghệ có độ tương đồng >= 85% mà không bị phụ thuộc vào các lỗi đánh máy hay ký tự đặc biệt."
    )

    doc.add_heading("2.3. Phân tích chức năng hệ thống:", level=2)
    p = doc.add_paragraph(
        "Hệ thống bao gồm các tác nhân chính là Khách vãng lai (Guest), Người dùng đã tải CV lên (User) và Hệ thống (Admin / Background Engine)."
    )

    doc.add_heading("2.3.1. Chức năng của guest:", level=3)
    p = doc.add_paragraph(
        "Khách vãng lai truy cập hệ thống có thể thực hiện:\n"
        "  - Xem bảng Salary Estimator: Tra cứu thông tin lương trung bình của các vai trò IT chính.\n"
        "  - Xem Dashboard Skill Market: Xem biểu đồ phân bổ nhu cầu tuyển dụng theo vai trò và khu vực địa lý dựa trên dữ liệu thực tế.\n"
        "  - Sử dụng JobRadar: Tìm kiếm liên kết công việc nhanh chóng trên các nền tảng việc làm lớn."
    )

    doc.add_heading("2.3.2. Chức năng của user:", level=3)
    p = doc.add_paragraph(
        "Người dùng sau khi tải lên tệp tin CV (.pdf hoặc .docx) sẽ được cấp quyền truy cập toàn bộ các chức năng chuyên sâu:\n"
        "  - Phân tích ATS Tune-Up: Xem điểm đánh giá ATS tĩnh cục bộ và nhận phản hồi chi tiết từ AI.\n"
        "  - So khớp CV-JD (JobMatcher): Nhập JD để nhận điểm số tương hợp, danh sách kỹ năng khớp và kỹ năng còn thiếu.\n"
        "  - Đề xuất việc làm (CareerMatch): Nhận danh sách 5 công việc tương thích nhất từ cơ sở dữ liệu.\n"
        "  - Lộ trình học tập (SkillBridge): Xem biểu đồ mạng nhện so sánh năng lực và nhận lộ trình học tập do AI đề xuất.\n"
        "  - Biên soạn CV (ResumeBuilder): Nhập thông tin biểu mẫu, nâng cấp mô tả dự án bằng AI theo chuẩn STAR và tải xuống file DOCX.\n"
        "  - Viết thư ứng tuyển (Cover Letter): Tạo Cover Letter tự động bám sát thế mạnh CV và yêu cầu của JD.\n"
        "  - Giả lập phỏng vấn (Interview Prep): Thực hành trả lời 5 câu hỏi phỏng vấn do AI sinh ra dựa trên CV/JD và nhận đánh giá chi tiết."
    )

    doc.add_heading("2.3.3. Chức năng của admin:", level=3)
    p = doc.add_paragraph(
        "Động cơ hệ thống hoạt động ở chế độ nền thực hiện các nhiệm vụ:\n"
        "  - Nạp và tiền xử lý cơ sở dữ liệu tuyển dụng IT từ tệp `JobsFE.csv`.\n"
        "  - Tự động mã hóa và lưu trữ ma trận nhúng vector việc làm `job_embeddings.npy` để đồng bộ chỉ mục FAISS.\n"
        "  - Duy trì, cập nhật từ điển kỹ năng chuẩn hóa cứng và mềm."
    )

    doc.add_heading("2.4. Biểu đồ Usecase và biểu đồ hoạt động:", level=2)
    
    doc.add_heading("2.4.1. Biểu đồ Usecase tổng quát:", level=3)
    p = doc.add_paragraph(
        "Biểu đồ Usecase tổng quát mô tả mối quan hệ tương tác giữa tác nhân Người dùng (Candidate/User) và Hệ thống CareerForge AI. "
        "Người dùng có thể tương tác với các ca sử dụng độc lập (JobRadar, ResumeBuilder) hoặc các ca sử dụng liên đới thông qua việc tải CV lên hệ thống "
        "(ATS Tune-Up, JobMatcher, CareerMatch, SkillBridge, CoverLetter, InterviewPrep)."
    )

    doc.add_heading("2.4.2. Biểu đồ hoạt động phân tích ATS (ATS Tune-Up):", level=3)
    p = doc.add_paragraph(
        "Luồng hoạt động phân tích ATS: Người dùng tải CV -> Hệ thống phân loại định dạng PDF/DOCX -> Trích xuất chuỗi ký tự thô -> "
        "Chạy đồng thời hai nhánh: 1) Nhánh Local quét 12 quy tắc tĩnh bằng Regex & spaCy để tính điểm số các tiêu chí; 2) Nhánh AI gửi văn bản thô "
        "sang Gemini API kèm cấu trúc Prompt tối ưu -> Tổng hợp kết quả và hiển thị so sánh Side-by-side trên giao diện."
    )

    doc.add_heading("2.4.3. Biểu đồ hoạt động so khớp CV-JD (JobMatcher):", level=3)
    p = doc.add_paragraph(
        "Luồng hoạt động so khớp: Nhập văn bản CV và JD -> spaCy trích xuất mảng kỹ năng CV và mảng kỹ năng JD -> Tính toán trọng số kỹ năng cứng dựa trên vị trí xuất hiện trong JD "
        "-> Tính toán độ bao phủ kỹ năng cứng và mềm -> Nhân hệ số trọng số phần (90% cứng + 10% mềm) -> Tính điểm tương thích tổng quan -> Hiển thị kỹ năng khớp/thiếu."
    )

    doc.add_heading("2.4.4. Biểu đồ hoạt động gợi ý việc làm (CareerMatch):", level=3)
    p = doc.add_paragraph(
        "Luồng gợi ý công việc: Trích xuất văn bản CV -> Khởi chạy SentenceTransformers tính vector nhúng CV -> Thực thi truy vấn FAISS Index để tìm Top N tin tuyển dụng "
        "có độ tương đồng ngữ nghĩa cao nhất -> Thực thi tính toán TF-IDF của CV với tin tuyển dụng để đo trùng khớp từ khóa -> Lai ghép điểm số Hybrid -> Sắp xếp và trả về Top 5."
    )

    doc.add_heading("2.4.5. Biểu đồ hoạt động cầu nối kỹ năng (SkillBridge):", level=3)
    p = doc.add_paragraph(
        "Luồng hoạt động: Người dùng chọn vai trò mục tiêu -> Hệ thống đọc kỹ năng yêu cầu chuẩn hóa từ `job_to_skill.json` -> So khớp với kỹ năng có sẵn trong CV "
        "-> Dựng biểu đồ Radar Chart so sánh tỷ lệ kỹ năng tương quan -> Đẩy danh sách kỹ năng thiếu sang Gemini API -> Nhận lộ trình tự học chi tiết theo tuần -> Hiển thị kết quả."
    )

    doc.add_heading("2.4.6. Biểu đồ hoạt động tạo và xuất bản CV (ResumeBuilder):", level=3)
    p = doc.add_paragraph(
        "Luồng tạo CV: Người dùng điền thông tin vào form -> Nhấp 'AI Enhance' để nâng cấp nội dung mô tả thô sang chuẩn STAR thông qua Gemini API -> "
        "Hệ thống tổng hợp dữ liệu JSON -> Gửi đến `generator_standard.py` để tạo tệp Word sử dụng phong cách theme đã chọn -> Xuất tệp `Resume.docx` về máy tính."
    )

    doc.add_heading("2.4.7. Biểu đồ hoạt động luyện tập phỏng vấn thử (Mock Interview):", level=3)
    p = doc.add_paragraph(
        "Luồng phỏng vấn: Gửi CV + JD sang Gemini API để sinh 5 câu hỏi phỏng vấn cá nhân hóa -> Hiển thị từng câu hỏi lên UI -> Người dùng nhập câu trả lời "
        "-> Gửi câu trả lời về Gemini API để chấm điểm, nhận xét ưu/nhược điểm và gợi ý câu trả lời mẫu -> Hiển thị kết quả đánh giá chi tiết."
    )

    doc.add_heading("2.5. Biểu đồ tuần tự:", level=2)
    
    doc.add_heading("2.5.1. Biểu đồ tuần tự chức năng phân tích ATS:", level=3)
    p = doc.add_paragraph(
        "Đặc tả luồng tuần tự và đặc tả API Endpoint `/api/ats-tuneup`:\n"
        "  - HTTP Method: POST\n"
        "  - Request Payload (Multipart/form-data): file (UploadFile), gemini_key (string), analysis_type (string)\n"
        "  - Response Schema (JSON):\n"
        "    { \"success\": bool, \"type\": \"local\"|\"ai\", \"results\": [...], \"scores\": { \"Ngữ pháp & Phong cách\": int, \"Thông tin Liên hệ\": int, \"Mật độ Từ khóa Kỹ năng\": int, \"Sự đa dạng Động từ\": int, \"Định dạng\": int } }\n"
        "Luồng đi: Streamlit UI -> gửi POST đến FastAPI `/api/ats-tuneup` -> Kích hoạt local resume_analysis hoặc ai analysis_enhancer -> Phản hồi kết quả JSON về UI."
    )

    doc.add_heading("2.5.2. Biểu đồ tuần tự chức năng so khớp và gợi ý việc làm:", level=3)
    p = doc.add_paragraph(
        "Đặc tả luồng tuần tự và đặc tả API Endpoint `/api/recommend-jobs`:\n"
        "  - HTTP Method: POST\n"
        "  - Request Payload (Multipart/form-data): file (UploadFile), search_method (string), top_n (int)\n"
        "  - Response Schema (JSON):\n"
        "    { \"success\": bool, \"personal_info\": { \"name\": string, \"email\": string, \"phone\": string, \"degree\": string, \"skills\": list }, \"recommended_jobs\": [ { \"position\": string, \"workplace\": string, \"salary\": float, \"job_role_and_duties\": string, \"requisite_skill\": string }, ... ] }\n"
        "Luồng đi: Streamlit UI gửi file CV -> FastAPI `/api/recommend-jobs` -> Gọi NLP Parser trích xuất text -> Gọi AI Recommender Engine tính toán độ trùng khớp ngữ nghĩa FAISS & TF-IDF -> Trả kết quả JSON về UI."
    )

    doc.add_heading("2.5.3. Biểu đồ tuần tự chức năng luyện phỏng vấn thử:", level=3)
    p = doc.add_paragraph(
        "Đặc tả luồng tuần tự và đặc tả API Endpoint `/api/interview/evaluate`:\n"
        "  - HTTP Method: POST\n"
        "  - Request Payload (JSON Application/json):\n"
        "    { \"resume_text\": string, \"jd_text\": string, \"questions\": list[string], \"answers\": list[string], \"gemini_key\": string }\n"
        "  - Response Schema (JSON):\n"
        "    { \"success\": bool, \"evaluation\": [ { \"question\": string, \"score\": int, \"positives\": list, \"negatives\": list, \"sample_answer\": string }, ... ] }\n"
        "Luồng đi: Streamlit UI gửi danh sách câu trả lời của ứng viên -> FastAPI `/api/interview/evaluate` -> Gọi Gemini API chấm điểm -> Trả về mảng JSON kết quả đánh giá chi tiết từng câu."
    )

    doc.add_heading("2.6. Tổng quan về cơ sở dữ liệu:", level=2)
    p = doc.add_paragraph(
        "Hệ thống CareerForge AI sử dụng cơ sở dữ liệu dạng phẳng (Flat File Database) kết hợp cơ sở dữ liệu vector để đảm bảo tính gọn nhẹ, bảo mật và tốc độ "
        "truy vấn nhanh chóng. Các thành phần CSDL bao gồm:"
    )
    
    db_items = [
        ("skills.json (Từ điển kỹ năng cứng)", 
         "- canonical_name: Tên kỹ năng chuẩn hóa (string, ví dụ: 'React')\n"
         "- aliases: Các từ khóa đồng nghĩa liên quan (array of strings, ví dụ: ['react', 'reactjs', 'react.js', 'react-js'])"),
        ("soft_skills.json (Từ điển kỹ năng mềm)", 
         "- canonical_name: Tên kỹ năng mềm chuẩn hóa (string, ví dụ: 'Communication')\n"
         "- aliases: Các từ đồng nghĩa liên quan (array of strings, ví dụ: ['communication', 'giao tiếp', 'thuyết trình'])"),
        ("job_to_skill.json (Bản đồ vai trò - kỹ năng)", 
         "- job_role: Tên vai trò công việc làm khóa (string, ví dụ: 'DevOps Engineer')\n"
         "- required_skills: Các kỹ năng cứng cốt lõi yêu cầu (array of strings, ví dụ: ['Docker', 'Kubernetes', 'Linux', 'CI/CD'])"),
        ("job_definition.json (Mô tả nghề nghiệp)", 
         "- job_role: Tên vai trò công việc làm khóa (string)\n"
         "- definition: Chuỗi văn bản mô tả chi tiết khái niệm và nhiệm vụ của vai trò đó (string)"),
        ("JobsFE.csv (Cơ sở dữ liệu tin tuyển dụng)", 
         "- id: Mã số định danh của tin tuyển dụng (int)\n"
         "- workplace: Địa điểm làm việc (string)\n"
         "- working_mode: Hình thức làm việc (string, ví dụ: 'full-time')\n"
         "- salary: Mức lương đề xuất (float)\n"
         "- position: Vị trí tuyển dụng (string)\n"
         "- job_role_and_duties: Mô tả nhiệm vụ chi tiết (string)\n"
         "- requisite_skill: Yêu cầu kỹ năng công nghệ (string)"),
        ("job_embeddings.npy (Chỉ mục Vector ngữ nghĩa)", 
         "- matrix: Ma trận Numpy kích thước N x 384 lưu trữ vector nhúng đặc trưng ngữ nghĩa của N tin tuyển dụng. Mỗi dòng biểu diễn một vector số thực 32-bit (float32) được tạo ra bởi mô hình SentenceTransformer.")
    ]
    
    for db_title, db_fields in db_items:
        p_db = doc.add_paragraph()
        p_db.paragraph_format.left_indent = Inches(0.3)
        p_db.add_run(f"• {db_title}:\n").bold = True
        p_db.add_run(db_fields)

    doc.add_page_break()

    # --- CHƯƠNG 3 ---
    doc.add_heading("CHƯƠNG 3: CÔNG NGHỆ SỬ DỤNG", level=1)
    
    doc.add_heading("3.1. Công nghệ phía frontend:", level=2)
    
    doc.add_heading("3.1.1. Nền tảng và điều hướng:", level=3)
    p = doc.add_paragraph(
        "Streamlit (phiên bản 1.45.1) được lựa chọn làm nền tảng xây dựng giao diện người dùng. Khác với các mô hình truyền thống "
        "yêu cầu phát triển Frontend (HTML/CSS/JS) độc lập với Backend, Streamlit cho phép lập trình toàn bộ giao diện trực tiếp "
        "bằng mã Python. Hệ thống thiết lập cấu trúc đa trang (multi-page app) bằng cách phân tách mã nguồn giao diện thành các tệp tin "
        "riêng biệt trong thư mục `pages/`. Cơ chế định tuyến (routing) của Streamlit tự động nhận diện và cập nhật thanh điều hướng sidebar "
        "dựa trên danh sách các tệp tin này, giúp đơn giản hóa đáng kể kiến trúc hệ thống."
    )

    doc.add_heading("3.1.2. Quản lý trạng thái và dữ liệu:", level=3)
    p = doc.add_paragraph(
        "Do Streamlit hoạt động theo cơ chế thực thi lại toàn bộ mã nguồn (re-run) mỗi khi có tương tác từ người dùng (nhấp nút, nhập text), "
        "việc duy trì trạng thái dữ liệu là một thách thức lớn. CareerForge AI sử dụng đối tượng `st.session_state` để lưu trữ các biến phiên "
        "như: Khóa API Gemini của người dùng, văn bản thô của CV sau khi tải lên, danh sách câu hỏi phỏng vấn được sinh ra, và các lịch sử chấm điểm. "
        "Điều này ngăn chặn việc hệ thống phải thực hiện parse lại CV hoặc gọi lại API Gemini trùng lặp, giúp giảm thiểu đáng kể chi phí băng thông và độ trễ."
    )

    doc.add_heading("3.1.3. Thiết kế giao diện:", level=3)
    p = doc.add_paragraph(
        "Nhằm mang lại trải nghiệm người dùng hiện đại, xóa bỏ giao diện thô sơ mặc định của Streamlit, dự án đã nhúng các đoạn mã Custom CSS chuyên sâu. "
        "Hệ thống áp dụng phong cách thiết kế kính mờ (Glassmorphism) sang trọng bằng cách thiết lập nền gradient mờ và hiệu ứng làm mờ hậu cảnh "
        "(`backdrop-filter: blur(8px)`). Bảng điều khiển trang chủ được chia cột linh hoạt bằng CSS Grid (`display: grid; grid-template-columns: repeat(3, 1fr)`) "
        "kết hợp các thẻ số liệu (metric cards) có hiệu ứng chuyển động rê chuột (hover transitions) nổi bật."
    )

    doc.add_heading("3.1.4. Các tiện ích bổ sung:", level=3)
    p = doc.add_paragraph(
        "Thư viện Pandas được tích hợp để thực hiện các thao tác tiền xử lý, lọc và cấu trúc hóa tập dữ liệu thị trường IT. "
        "Để biểu diễn số liệu trực quan, dự án sử dụng Altair - thư viện đồ họa khai báo cho Python. Altair giúp xây dựng biểu đồ thanh ngang "
        "thể hiện nhu cầu việc làm và biểu đồ Donut phân bổ địa lý tương tác cao, hỗ trợ tự động hiển thị tooltip thông tin chi tiết khi người dùng "
        "rê chuột qua các phân đoạn biểu đồ."
    )

    doc.add_heading("3.2. Công nghệ phía backend và database:", level=2)
    
    doc.add_heading("3.2.1. Nền tảng và framework:", level=3)
    p = doc.add_paragraph(
        "Tầng Backend của ứng dụng được phân tách độc lập và đóng gói dưới dạng dịch vụ REST API sử dụng FastAPI. "
        "FastAPI hoạt động trên nền tảng ASGI (Asynchronous Server Gateway Interface) bằng Uvicorn server, mang lại hiệu năng xử lý bất đồng bộ "
        "cực cao. Việc sử dụng FastAPI giúp mở rộng khả năng tích hợp của hệ thống, cho phép các nền tảng di động hoặc ứng dụng web bên ngoài "
        "gọi các dịch vụ phân tích CV, so khớp và gợi ý việc làm."
    )
    p = doc.add_paragraph(
        "Hệ thống kết hợp các mô hình toán học và dịch vụ AI đám mây để tính toán độ tương hợp ngữ nghĩa việc làm:"
    )
    
    p_st = doc.add_paragraph()
    p_st.paragraph_format.left_indent = Inches(0.3)
    p_st.add_run("• Mô hình SentenceTransformers (paraphrase-MiniLM-L6-v2):\n").bold = True
    p_st.add_run(
        "  Sử dụng mô hình Transformer nhỏ gọn và hiệu quả để mã hóa văn bản tự nhiên thành các vector nhúng (dense vector) 384 chiều. "
        "Mô hình được lượng tử hóa động (dynamic quantization) sang dạng số nguyên 8-bit (qint8) bằng PyTorch nhằm giảm kích thước bộ nhớ "
        "và tăng tốc độ tính toán vector trên CPU gấp 2 lần. Không gian vector 384 chiều này đại diện cho ý nghĩa ngữ nghĩa của văn bản: "
        "các câu văn hoặc CV/JD có ý nghĩa tương đương nhau sẽ có khoảng cách vector gần nhau hơn trong không gian, khắc phục hoàn toàn nhược điểm "
        "của phương pháp so khớp từ khóa chính xác truyền thống."
    )
    
    p_faiss = doc.add_paragraph()
    p_faiss.paragraph_format.left_indent = Inches(0.3)
    p_faiss.add_run("• Cơ sở dữ liệu Vector FAISS (IndexFlatIP):\n").bold = True
    p_faiss.add_run(
        "  Sử dụng thư viện FAISS để lập chỉ mục và thực hiện tìm kiếm láng giềng gần nhất (k-Nearest Neighbors) trên CPU cực nhanh. "
        "Chỉ mục sử dụng là IndexFlatIP (Inner Product). Đối với các vector nhúng đã được chuẩn hóa độ dài L2 (L2 normalized vectors), "
        "phép tính Inner Product (tích vô hướng) tương đương hoàn toàn với độ tương đồng Cosine (Cosine Similarity). "
        "Công thức tính Cosine Similarity giữa vector CV (A) và vector JD (B) được biểu diễn:\n"
        "  Cosine Similarity(A, B) = (A • B) / (||A|| * ||B||)\n"
        "  Với giá trị nằm trong khoảng [-1, 1], càng gần 1 thể hiện mức độ tương hợp ngữ nghĩa càng cao."
    )

    p_tfidf = doc.add_paragraph()
    p_tfidf.paragraph_format.left_indent = Inches(0.3)
    p_tfidf.add_run("• Trọng số từ khóa TF-IDF (Term Frequency - Inverse Document Frequency):\n").bold = True
    p_tfidf.add_run(
        "  Đo lường tầm quan trọng của một từ khóa kỹ năng trong tập hợp văn bản. "
        "TF (Term Frequency) tính tần suất xuất hiện của từ trong tài liệu. IDF (Inverse Document Frequency) giảm trọng số của các từ xuất hiện phổ biến "
        "ở mọi tài liệu (như 'the', 'and', 'developer') và tăng trọng số cho các từ khóa công nghệ đặc trưng (như 'Kubernetes', 'FastAPI')."
    )

    p_hybrid = doc.add_paragraph()
    p_hybrid.paragraph_format.left_indent = Inches(0.3)
    p_hybrid.add_run("• Thuật toán tìm kiếm lai ghép (Hybrid Search):\n").bold = True
    p_hybrid.add_run(
        "  Kết hợp sức mạnh của tìm kiếm ngữ nghĩa sâu (FAISS) và tìm kiếm từ khóa chính xác (TF-IDF) nhằm tránh hiện tượng bỏ sót các từ khóa công nghệ bắt buộc "
        "hoặc nhận diện nhầm các ngữ cảnh không liên quan. Cả hai điểm số FAISS và TF-IDF được chuẩn hóa Min-Max về khoảng [0, 1] theo công thức:\n"
        "  Normalized_Score = (S - S_min) / (S_max - S_min)\n"
        "  Sau đó, điểm số Hybrid tổng hợp được tính theo công thức tuyến tính kết hợp trọng số alpha (mặc định alpha = 0.5):\n"
        "  S_hybrid = alpha * S_faiss + (1 - alpha) * S_tfidf\n"
        "  Nhờ sự lai ghép này, kết quả gợi ý việc làm vừa đảm bảo đúng định hướng ngành nghề vừa khớp các từ khóa công nghệ bắt buộc."
    )

    p_gemini = doc.add_paragraph()
    p_gemini.paragraph_format.left_indent = Inches(0.3)
    p_gemini.add_run("• Google Gemini API & Structured Output JSON Schema:\n").bold = True
    p_gemini.add_run(
        "  Tích hợp mô hình Gemini 1.5 Flash thông qua các thiết kế Prompt kỹ lưỡng. Hệ thống sử dụng tính năng 'Structured Output' bằng cách cung cấp "
        "JSON schema cụ thể để bắt buộc mô hình trả về dữ liệu đúng định dạng mong muốn (ví dụ: mảng JSON chứa các câu hỏi phỏng vấn hoặc bảng kết quả chấm điểm). "
        "Điều này giúp triệt tiêu hiện tượng mô hình sinh từ ngữ thừa (hallucination) và đảm bảo tính ổn định khi parse dữ liệu từ FastAPI Backend."
    )

    doc.add_heading("3.2.2. Database:", level=3)
    p = doc.add_paragraph(
        "Hệ thống lưu trữ cơ sở dữ liệu vector thông qua chỉ mục FAISS trên CPU. Tệp nhị phân `job_embeddings.npy` chứa ma trận vector biểu diễn ngữ nghĩa "
        "của 7.000 tin tuyển dụng được nạp vào bộ nhớ RAM ngay khi máy chủ FastAPI khởi tạo. Các dữ liệu quan hệ kỹ năng được lưu trữ gọn nhẹ dưới dạng "
        "các tệp tin phẳng JSON giúp tối ưu tốc độ đọc ghi, không tốn tài nguyên quản lý kết nối cơ sở dữ liệu."
    )

    doc.add_heading("3.2.3. Bảo mật và phân quyền:", level=3)
    p = doc.add_paragraph(
        "Để đảm bảo an toàn thông tin, FastAPI được cấu hình CORS Middleware chặn các cuộc gọi trái phép hoặc cấp quyền gọi liên miền phục vụ tích hợp. "
        "Đặc biệt, hệ thống áp dụng cơ chế xử lý tệp tin hoàn toàn trong bộ nhớ đệm (In-memory file processing). Khi người dùng tải CV lên, "
        "dữ liệu nhị phân được đọc và truyền trực tiếp dưới dạng luồng byte (`BytesIO`) đi thẳng vào bộ parser PyMuPDF/python-docx nằm trong bộ nhớ RAM, "
        "hoàn toàn không ghi bất kỳ tệp tin vật lý nào lên đĩa cứng của máy chủ. Điều này ngăn chặn triệt để rủi ro rò rỉ dữ liệu cá nhân (Data Leakage) "
        "và bảo vệ thông tin riêng tư của ứng viên."
    )

    doc.add_heading("3.3. Mã nguồn:", level=2)
    p = doc.add_paragraph(
        "Cấu trúc mã nguồn của dự án được module hóa rõ ràng, phân tách mã giao diện người dùng (Presentation) và mã xử lý logic (Business Logic):"
    )
    
    # Code Architecture tree
    p_code = doc.add_paragraph()
    p_code.paragraph_format.line_spacing = 1.15
    p_code.paragraph_format.space_after = Pt(2)
    
    code_structure = (
        "opportune-jobMate-1/\n"
        "├── Home.py                    # Điểm chạy giao diện chính Streamlit\n"
        "├── requirements.txt           # Danh sách thư viện\n"
        "├── api/\n"
        "│   └── main.py                # Điểm chạy API FastAPI Server\n"
        "├── analyzer/\n"
        "│   ├── resume_analysis.py     # Logic 12 kiểm tra ATS cục bộ\n"
        "│   └── analysis_enhancer.py   # AI phân tích CV nâng cao\n"
        "├── preprocessor/\n"
        "│   ├── parser.py              # Đọc tệp PDF/DOCX sang text\n"
        "│   ├── skills.py              # Trích xuất kỹ năng n-gram và RapidFuzz\n"
        "│   ├── personal_info.py       # Regex thông tin cá nhân\n"
        "│   └── jd_section_parser.py   # Phân đoạn JD thô\n"
        "├── recommender/\n"
        "│   ├── ai_model.py            # FAISS, SentenceTransformers & Hybrid Search\n"
        "│   └── top_n_jobs.py          # So khớp kỹ năng cơ bản\n"
        "├── builder/\n"
        "│   ├── generator_standard.py  # Tạo file Word (.docx) theo theme\n"
        "│   └── resume_enhancer.py     # AI nâng cấp câu viết chuẩn STAR\n"
        "├── pages/\n"
        "│   ├── 1_📡_JobRadar.py        # Trang tìm việc đa cổng\n"
        "│   ├── 2_🔎_JobMatcher.py      # Trang chấm điểm CV-JD\n"
        "│   ├── 3_💼_CareerMatch.py     # Gợi ý công việc FAISS\n"
        "│   ├── 4_📚_SkillBridge.py     # Radar map kỹ năng thiếu\n"
        "│   ├── 5_📝_ResumeBuilder.py   # Biên soạn CV động\n"
        "│   ├── 6_🛠️_ATS_TuneUp.py      # Đánh giá ATS side-by-side\n"
        "│   ├── 7_✉️_Cover_Letter.py    # Sinh thư ứng tuyển AI\n"
        "│   └── 8_🎙️_Interview_Prep.py   # Giả lập phỏng vấn thử\n"
        "└── ui/\n"
        "    ├── style.css              # Custom Glassmorphism CSS\n"
        "    └── render_footer.py       # Kết xuất chân trang"
    )
    p_code.add_run(code_structure).font.name = 'Consolas'
    p_code.runs[0].font.size = Pt(9.5)
    
    doc.add_paragraph() 

    doc.add_page_break()

    # --- CHƯƠNG 4 ---
    doc.add_heading("CHƯƠNG 4: TRIỂN KHAI VÀ THỰC NGHIỆM", level=1)
    
    doc.add_heading("4.1. Giao diện trang chủ và Dashboard Skill Market:", level=2)
    p = doc.add_paragraph(
        "Trang chủ `Home.py` chào đón người dùng bằng một giao diện Dashboard trực quan. Phía trên hiển thị bảng trạng thái hệ thống cho biết động cơ NLP spaCy "
        "và chỉ mục ngữ nghĩa FAISS đã được nạp thành công. Biểu đồ tròn Donut phân tích địa lý việc làm cho thấy tỷ lệ tuyển dụng tập trung "
        "chủ yếu ở TP. Hồ Chí Minh (52%) và Hà Nội (38%). Bảng Salary Estimator tương tác cho phép người dùng chọn nhanh vai trò (ví dụ: Backend Developer) "
        "và cấp độ (ví dụ: Mid-level) để xem thống kê lương dao động 1.200 - 2.300 USD kèm mức độ cạnh tranh hồ sơ Cao."
    )

    doc.add_heading("4.2. Giao diện tìm kiếm và định hướng việc làm (JobRadar):", level=2)
    p = doc.add_paragraph(
        "Giao diện trang `1_📡_JobRadar.py` cung cấp các trường nhập liệu từ khóa. Khi ứng viên nhập từ khóa vai trò mục tiêu và nhấp tìm kiếm, "
        "hệ thống tự động sinh ra các UI card đại diện cho các nền tảng việc làm kèm biểu tượng sinh động. Người dùng nhấp chọn sẽ kích hoạt trình duyệt "
        "mở tab mới đi thẳng vào kết quả tìm kiếm thời gian thực của LinkedIn, Indeed hay VietnamWorks tương ứng."
    )

    doc.add_heading("4.3. Giao diện so khớp CV với mô tả công việc (JobMatcher):", level=2)
    p = doc.add_paragraph(
        "Trang `2_🔎_JobMatcher.py` hỗ trợ tải lên tệp tin CV và hộp văn bản dán nội dung JD tuyển dụng. Khi kích hoạt so khớp, hệ thống hiển thị kết quả chấm điểm "
        "tổng quan dạng phần trăm (%). Dưới điểm số là hai cột so sánh chi tiết: 1) Kỹ năng chuyên môn trùng khớp (Matched hard skills) và kỹ năng chuyên môn bị thiếu "
        "(Missing hard skills); 2) Kỹ năng mềm trùng khớp và kỹ năng mềm bị thiếu. Các kỹ năng thiếu được đính kèm gợi ý tài nguyên học tập hữu ích."
    )

    doc.add_heading("4.4. Giao diện kết quả phân tích kỹ năng trùng và thiếu từ JobMatcher:", level=2)
    p = doc.add_paragraph(
        "Màn hình kết quả hiển thị hai khối thông tin song song: Khối kỹ năng khớp (màu xanh lá) giúp ứng viên biết điểm mạnh hiện tại, "
        "và khối kỹ năng thiếu (màu cam nhạt) cảnh báo các yêu cầu kỹ năng trong JD mà CV chưa thể hiện được. Đây là cơ sở để ứng viên cập nhật "
        "bổ sung từ khóa tương ứng trước khi nộp hồ sơ."
    )

    doc.add_heading("4.5. Giao diện gợi ý nghề nghiệp tự động (CareerMatch):", level=2)
    p = doc.add_paragraph(
        "Trang `3_💼_CareerMatch.py` cho phép tải lên hồ sơ cá nhân. Hệ thống tự động kích hoạt thuật toán Hybrid Search để đối sánh CV với cơ sở dữ liệu. "
        "Màn hình hiển thị danh sách Top 5 vai trò IT phù hợp nhất. Mỗi vai trò được trình bày trong một hộp Glassmorphism card thể hiện: Tên vai trò, "
        "số kỹ năng trùng khớp, danh sách các kỹ năng của ứng viên tương thích với vai trò và mô tả nhiệm vụ công việc rút trích từ dữ liệu hệ thống."
    )

    doc.add_heading("4.6. Giao diện cầu nối kỹ năng và biểu đồ mạng nhện (SkillBridge):", level=2)
    p = doc.add_paragraph(
        "Trang `4_📚_SkillBridge.py` hiển thị một biểu đồ mạng nhện Radar Chart trực quan được dựng từ thư viện đồ họa. Biểu đồ so sánh trực quan tỷ lệ kỹ năng "
        "ứng viên đang có so với bộ kỹ năng yêu cầu chuẩn hóa của một vai trò do ứng viên lựa chọn, giúp phát hiện trực quan khoảng trống kỹ năng."
    )

    doc.add_heading("4.7. Giao diện lộ trình tự học đề xuất bởi AI từ SkillBridge:", level=2)
    p = doc.add_paragraph(
        "Phía dưới biểu đồ Radar, hệ thống hiển thị bảng lộ trình tự học (learning roadmap) chi tiết theo từng tuần cho ứng viên. Lộ trình được sinh tự động "
        "bởi Gemini API, chỉ rõ tài liệu học tập, các dự án nhỏ cần thực hiện để bù đắp các công nghệ còn thiếu."
    )

    doc.add_heading("4.8. Giao diện biểu mẫu động tạo CV (ResumeBuilder):", level=2)
    p = doc.add_paragraph(
        "Trang `5_📝_ResumeBuilder.py` hiển thị giao diện nhập liệu dạng tab ngăn nắp (Thông tin cá nhân, Học vấn, Kinh nghiệm, Dự án...). Người dùng có thể "
        "nhấp chọn dấu cộng (+) để chèn thêm các dòng mô tả công việc hoặc dự án mới, giúp biểu mẫu co giãn động theo nội dung thực tế."
    )

    doc.add_heading("4.9. Giao diện tối ưu hóa mô tả công việc theo cấu trúc STAR bằng AI:", level=2)
    p = doc.add_paragraph(
        "Đối với mỗi dòng mô tả kinh nghiệm hoặc dự án trong ResumeBuilder, nút bấm 'AI Enhance' có sẵn giúp người dùng "
        "chuyển đổi câu văn thô sơ ban đầu thành các mô tả dự án chuyên nghiệp chuẩn cấu trúc STAR (Situation, Task, Action, Result) bằng cách "
        "bổ sung các động từ mạnh mẽ và con số định lượng thành tựu thông qua Gemini API."
    )

    doc.add_heading("4.10. Giao diện tải xuống tệp CV Word (.docx) chuẩn ATS:", level=2)
    p = doc.add_paragraph(
        "Người dùng sau khi tối ưu thông tin nhấp chọn tải xuống tài liệu. Hệ thống gọi module `generator_standard.py` tạo tệp Word trong bộ nhớ đệm "
        "và tải xuống máy tính với tên mặc định `Resume.docx`. Định dạng tệp Word sạch sẽ, font Cambria sắc nét, lề narrow đều đặn, "
        "và các mốc thời gian được căn lề phải thẳng hàng thông qua kỹ thuật thiết lập Tab Stops của python-docx, đảm bảo thân thiện tối đa với ATS."
    )

    doc.add_heading("4.11. Giao diện đánh giá ATS tĩnh (ATS Tune-Up Local):", level=2)
    p = doc.add_paragraph(
        "Trang `6_🛠️_ATS_TuneUp.py` cho phép người dùng xem kết quả đánh giá hồ sơ. Cột bên trái hiển thị kết quả quét 12 tiêu chí local "
        "dưới dạng điểm số danh mục (%) và các thẻ Warnings (màu cam) hoặc Successes (màu xanh). Đây là các kiểm tra kỹ thuật tĩnh giúp phát hiện "
        "nhanh lỗi định dạng hoặc lạm dụng từ ngữ."
    )

    doc.add_heading("4.12. Giao diện phân tích chất lượng CV bằng AI (ATS Tune-Up AI):", level=2)
    p = doc.add_paragraph(
        "Cột bên phải hiển thị đánh giá định tính sâu sắc của Gemini AI chỉ ra các lỗi hành văn, lỗi ngữ pháp thực tế và khuyến nghị viết lại trực tiếp "
        "cho từng phần cụ thể của CV, bổ sung hoàn hảo cho bộ quét tĩnh cục bộ."
    )

    doc.add_heading("4.13. Giao diện soạn thảo Cover Letter cá nhân hóa bằng AI:", level=2)
    p = doc.add_paragraph(
        "Trang `7_✉️_Cover_Letter.py` cho phép người dùng chọn tông giọng (Professional, Confident, Creative). Sau khi chạy, Cover Letter được sinh ra "
        "và hiển thị trong một khung soạn thảo văn bản. Người dùng có thể chỉnh sửa trực tiếp nội dung bức thư ngay trên trang web."
    )

    doc.add_heading("4.14. Giao diện tải xuống tệp Cover Letter (.docx) trang trọng:", level=2)
    p = doc.add_paragraph(
        "Người dùng nhấp nút 'Tải xuống DOCX' tại trang Cover Letter để nhận tệp Word. Hệ thống tự động thiết kế tiêu đề thư chứa tên, email, "
        "điện thoại của ứng viên trình bày trang trọng ở đầu trang, căn lề văn bản cân đối, sẵn sàng để gửi kèm hồ sơ xin việc."
    )

    doc.add_heading("4.15. Giao diện giả lập phòng phỏng vấn thử (Mock Interview):", level=2)
    p = doc.add_paragraph(
        "Giao diện trang `8_🎙️_Interview_Prep.py` mô phỏng một phòng phỏng vấn ảo. Hệ thống hiển thị 5 câu hỏi phỏng vấn được thiết kế riêng "
        "dựa trên CV của ứng viên và JD ứng tuyển. Bên dưới mỗi câu hỏi là một ô nhập liệu văn bản lớn để ứng viên nhập câu trả lời."
    )

    doc.add_heading("4.16. Giao diện kết quả chấm điểm phỏng vấn tổng quan từ AI:", level=2)
    p = doc.add_paragraph(
        "Sau khi ứng viên hoàn thành 5 câu trả lời và nhấp nộp bài, giao diện phỏng vấn sẽ chuyển sang trang kết quả. Phía trên hiển thị điểm số tổng quan "
        "của buổi phỏng vấn (thang điểm 10) cùng nhận xét chung về thái độ, kỹ năng diễn đạt và mức độ phù hợp của ứng viên."
    )

    doc.add_heading("4.17. Giao diện phân tích câu trả lời chi tiết và câu trả lời mẫu từ AI:", level=2)
    p = doc.add_paragraph(
        "Phía dưới điểm tổng quan, hệ thống liệt kê chi tiết từng câu hỏi: Điểm số của câu, nhận xét điểm tốt đạt được (Positives), "
        "các ý bị thiếu hoặc lỗi diễn đạt (Negatives), và một câu trả lời mẫu tối ưu (Sample Answer) bằng chữ in nghiêng màu xám để ứng viên học hỏi cách trả lời chuẩn."
    )

    doc.add_heading("4.18. Giao diện tài liệu API tự động (FastAPI Swagger UI):", level=2)
    p = doc.add_paragraph(
        "Hệ thống FastAPI tự động phát sinh trang tài liệu API tại đường dẫn `/docs` sử dụng Swagger UI. Giao diện hiển thị danh sách toàn bộ các endpoint "
        "REST API của hệ thống (như `/api/ats-tuneup`, `/api/recommend-jobs`, `/api/interview/evaluate`), cho phép lập trình viên kiểm thử trực tiếp "
        "các request/response một cách trực quan."
    )

    doc.add_heading("4.19. Những hạn chế và hướng phát triển của dự án:", level=2)
    p = doc.add_paragraph(
        "Đồ án đã đạt được các kết quả thực tế khả quan nhưng vẫn tồn tại một số điểm hạn chế kỹ thuật: "
        "1) Cơ sở dữ liệu và từ điển kỹ năng hiện tại mới chỉ tối ưu tốt cho ngành Công nghệ thông tin (IT), chưa mở rộng sang các khối ngành khác như Kinh tế, Marketing; "
        "2) Chức năng Mock Interview mới chỉ hỗ trợ nhập câu trả lời qua bàn phím văn bản, chưa hỗ trợ giao tiếp giọng nói."
    )
    p = doc.add_paragraph(
        "Hướng phát triển tương lai:\n"
        "• Tích hợp công nghệ Speech-to-Text (như Whisper của OpenAI) và Text-to-Speech để ứng viên có thể thực hành trả lời bằng giọng nói trực tiếp, "
        "giúp giả lập không khí phòng phỏng vấn chân thực hơn.\n"
        "• Mở rộng quy mô cơ sở dữ liệu tin tuyển dụng và từ điển kỹ năng sang các nhóm ngành nghề khác ngoài IT để phục vụ đa dạng đối tượng người dùng."
    )

    # --- REFERENCES ---
    doc.add_page_break()
    doc.add_heading("TÀI LIỆU THAM KHẢO", level=1)
    doc.add_paragraph("1. Ths. Vũ Văn Quang, Hướng dẫn môn học Đồ Án Cơ Sở, Trường Công nghệ Thông tin Phenikaa, 2026.", style='List Bullet')
    doc.add_paragraph("2. Google Generative AI Documentation, Gemini API Developer Guide: Prompt Engineering & Structured JSON Output, 2025.", style='List Bullet')
    doc.add_paragraph("3. spaCy NLP Library Documentation, Linguistic Features: Tokenization, POS Tagging, NER, and N-Grams pipeline, 2025.", style='List Bullet')
    doc.add_paragraph("4. Facebook AI Similarity Search (FAISS) GitHub Repository, Indexing and Similarity Search for Dense Vectors, 2025.", style='List Bullet')
    doc.add_paragraph("5. python-docx Community, Word Document Generation and Formatting with python-docx library, 2025.", style='List Bullet')

    # Save to file
    output_filename = "Bao_cao_Do_an_co_so_CareerForge_AI_Final.docx"
    doc.save(output_filename)
    print(f"SUCCESS: Report saved to {output_filename}")

if __name__ == "__main__":
    build_report()
