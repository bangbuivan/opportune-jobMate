import docx

def read_structure(file_path):
    doc = docx.Document(file_path)
    output = []
    output.append(f"Total paragraphs: {len(doc.paragraphs)}")
    output.append(f"Total tables: {len(doc.tables)}")
    
    headings_found = 0
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        
        is_heading_style = 'Heading' in p.style.name or p.style.name.startswith('Title') or p.style.name.startswith('Subtitle') or 'heading' in p.style.name.lower()
        is_numbered = any(text.startswith(prefix) for prefix in ["CHƯƠNG", "Chương", "Phần", "PHẦN", "MỤC LỤC"])
        is_numbered_sub = len(text) > 0 and text[0].isdigit() and ('.' in text[:5] or ' ' in text[:5])
        
        if is_heading_style or is_numbered or (is_numbered_sub and len(text) < 120):
            output.append(f"[{p.style.name}] Line {idx}: {text}")
            headings_found += 1

    with open("dacs_structure.txt", "w", encoding="utf-8") as f:
        f.write("\n".join(output))
    print("SUCCESS: Written structure to dacs_structure.txt")

if __name__ == "__main__":
    try:
        read_structure("DACS.docx")
    except Exception as e:
        print("Error:", e)
