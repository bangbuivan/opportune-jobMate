import docx

def read_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    
    # Read paragraphs
    for para in doc.paragraphs:
        full_text.append(para.text)
        
    # Read tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            full_text.append(" | ".join(row_text))
            
    return "\n".join(full_text)

if __name__ == "__main__":
    try:
        content = read_docx("plan.docx")
        with open("plan_content.txt", "w", encoding="utf-8") as f:
            f.write(content)
        print("Success! Written to plan_content.txt")
    except Exception as e:
        print(f"Error: {e}")
