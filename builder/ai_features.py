import google.generativeai as genai
import json
import re
import random
import spacy
from preprocessor.spacy_nlp import load_spacy_nlp_model
from preprocessor.skills import extract_skills_fuzzy
import preprocessor.personal_info as pf
from builder.resume_enhancer import get_suitable_gemini_model

# Load spaCy nlp model locally for parsing
try:
    nlp = load_spacy_nlp_model()
except Exception:
    nlp = None

def extract_job_details(jd_text):
    """Heuristically extract job title and company from JD text."""
    job_title = "Software Engineer"
    company = "your company"
    
    if not jd_text:
        return job_title, company
        
    # Search for Job Title patterns
    title_match = re.search(r"(?:job title|position|role|seeking a|hiring a)[:\-–\s]*([A-Za-z0-9\s#\+\.]{3,30})", jd_text, re.IGNORECASE)
    if title_match:
        job_title = title_match.group(1).strip()
    else:
        # Fallback to first line if it looks like a title
        first_line = jd_text.splitlines()[0].strip()
        if len(first_line) > 5 and len(first_line) < 50:
            job_title = first_line
            
    # Search for Company name patterns
    company_match = re.search(r"(?:company|organization|at|about)[:\-–\s]*([A-Za-z0-9\s\.\,]{3,30})", jd_text, re.IGNORECASE)
    if company_match:
        company = company_match.group(1).strip()
        
    return job_title, company

def generate_local_cover_letter(resume_text, jd_text, tone):
    """Generate a template-based cover letter using extracted info in Vietnamese."""
    global nlp
    if nlp is None:
        nlp = load_spacy_nlp_model()
        
    doc = nlp(resume_text)
    name = pf.extract_name(doc, resume_text)
    email = pf.extract_mail(resume_text) or "applicant@example.com"
    phone = pf.extract_phone(resume_text) or "555-0199"
    skills = extract_skills_fuzzy(doc)
    
    job_title, company = extract_job_details(jd_text)
    
    skills_str = ", ".join(skills[:5]) if skills else "phát triển phần mềm, giải quyết vấn đề và thiết kế hệ thống"
    
    # Heuristics based on tone (translated to Vietnamese)
    if tone == "Conversational" or tone == "Trò chuyện":
        intro = f"Kính gửi đội ngũ tuyển dụng tại {company},\n\nTôi viết thư này vì rất hào hứng với vị trí {job_title}. Xem qua bản mô tả công việc, tôi thấy đây là một sự kết hợp hoàn hảo cho những gì tôi đam mê thực hiện."
        body = f"Trong suốt sự nghiệp của mình, tôi đã tập trung xây dựng các hệ thống đáng tin cậy và cộng tác tốt với các đội ngũ xuất sắc. Tôi đã phát triển các kỹ năng vững chắc về {skills_str}, và tôi tin rằng chúng sẽ giúp tôi nhanh chóng thích nghi và đóng góp cho công việc."
        outro = f"Tôi rất mong có cơ hội trò chuyện nhiều hơn về cách tôi có thể đóng góp cho công ty. Hãy cho tôi biết nếu chúng ta có thể sắp xếp một buổi trao đổi!\n\nTrân trọng,\n{name}"
    elif tone == "Enthusiastic" or tone == "Nhiệt huyết":
        intro = f"Kính gửi Đội ngũ Tuyển dụng tại {company},\n\nTôi vô cùng hào hứng khi thấy thông tin tuyển dụng cho vị trí {job_title}! Tôi đã theo dõi hành trình của đội ngũ các bạn và thực sự mong muốn đóng góp vào sứ mệnh của công ty."
        body = f"Bộ kỹ năng kỹ thuật của tôi, bao gồm chuyên môn về {skills_str}, hoàn toàn phù hợp với các yêu cầu mà quý công ty đã nêu ra. Tôi mang đến nguồn năng lượng cao, niềm đam mê giải quyết các thử thách kỹ thuật phức tạp và cam kết hướng tới sự xuất sắc."
        outro = f"Tôi rất hy vọng có cơ hội mang nhiệt huyết này đến với đội ngũ của bạn. Xin chân thành cảm ơn thời gian và sự cân nhắc của các bạn!\n\nTrân trọng,\n{name}"
    elif tone == "Confident" or tone == "Tự tin":
        intro = f"Kính gửi Trưởng bộ phận Tuyển dụng,\n\nTôi viết thư này để chứng minh sự phù hợp mạnh mẽ của mình với vị trí {job_title} tại {company}. Với nền tảng và những thành tựu đã được chứng minh, tôi tự tin sẽ mang lại giá trị tức thì cho công ty."
        body = f"Kinh nghiệm của tôi trải dài trên nhiều lĩnh vực kỹ thuật khác nhau, được hỗ trợ bởi sự thành thạo thực tế về {skills_str}. Tôi xuất sắc trong việc chuyển dịch các yêu cầu phức tạp thành các giải pháp mạnh mẽ, hiệu suất cao và thúc đẩy hoàn thành các mục tiêu kỹ thuật."
        outro = f"Tôi đã sẵn sàng mang sự tận tụy và các kỹ năng của mình đến đóng góp cho {company}. Tôi rất mong được thảo luận chi tiết hơn về cách kinh nghiệm của mình đáp ứng nhu cầu của quý công ty.\n\nTrân trọng,\n{name}"
    else:  # Professional (Default)
        intro = f"Kính gửi Đội ngũ Tuyển dụng tại {company},\n\nTôi viết thư này để bày tỏ sự quan tâm sâu sắc của mình đối với vị trí {job_title}. Với nền tảng vững chắc về phát triển phần mềm cùng các năng lực phù hợp với yêu cầu tuyển dụng, tôi rất mong muốn được thảo luận về sự phù hợp của mình."
        body = f"Trong suốt hành trình phát triển chuyên môn, tôi đã trau dồi chuyên môn trong các lĩnh vực chính, bao gồm {skills_str}. Tôi tập trung vào việc bàn giao mã nguồn chất lượng cao, có khả năng mở rộng và phối hợp chặt chẽ với các đội ngũ liên chức năng để đạt được các mục tiêu của tổ chức."
        outro = f"Xin chân thành cảm ơn quý công ty đã cân nhắc hồ sơ của tôi. Tôi rất mong có cơ hội thảo luận về cách các năng lực của tôi có thể đóng góp vào sự phát triển của công ty.\n\nTrân trọng,\n{name}"
 
    letter = f"""{name}
{email} | {phone}
 
{intro}
 
{body}
 
{outro}"""
    return letter
 
def generate_local_interview_questions(resume_text, jd_text):
    """Generate 5 relevant questions using extracted skills in Vietnamese."""
    global nlp
    if nlp is None:
        nlp = load_spacy_nlp_model()
        
    doc = nlp(resume_text)
    skills = extract_skills_fuzzy(doc)
    
    # Common skills questions mapping with multiple questions per skill for diversity
    skills_questions = {
        "python": [
            "Sự khác biệt giữa list và tuple trong Python là gì, và khi nào bạn nên sử dụng generator?",
            "Trình bày cơ chế hoạt động của Decorator trong Python. Hãy nêu một ví dụ thực tế mà bạn đã áp dụng decorator để giải quyết vấn đề.",
            "Cơ chế quản lý bộ nhớ trong Python hoạt động như thế nào? Sự khác biệt giữa Reference Counting và Garbage Collection là gì?",
            "Giải thích khái niệm GIL (Global Interpreter Lock) trong Python. Nó ảnh hưởng như thế nào đến lập trình đa luồng (multi-threading) và cách khắc phục?",
            "Sự khác biệt giữa __new__ và __init__ trong Python là gì? Khi nào bạn cần ghi đè (override) phương thức __new__?"
        ],
        "java": [
            "Giải thích khái niệm quản lý bộ nhớ của JVM và cách hoạt động của bộ thu gom rác (garbage collection) trong Java.",
            "Sự khác biệt giữa Abstract Class và Interface trong Java là gì? Từ Java 8 trở đi, Interface có những điểm mới nào?",
            "Trình bày cơ chế xử lý ngoại lệ (Exception Handling) trong Java. Khác biệt giữa Checked Exception và Unchecked Exception là gì?",
            "Khái niệm ThreadPool trong Java hoạt động thế nào? Trình bày cách tối ưu số lượng thread trong một ThreadPool cho tác vụ CPU-bound và I/O-bound.",
            "Cơ chế Reflection trong Java là gì? Hãy nêu lợi ích cũng như rủi ro về mặt hiệu năng/bảo mật khi sử dụng Reflection."
        ],
        "sql": [
            "Làm thế nào để tối ưu hóa một truy vấn cơ sở dữ liệu bị chậm, và mục đích của việc đánh chỉ mục (indexing) là gì?",
            "Giải thích sự khác biệt giữa các mức độ cô lập giao dịch (Transaction Isolation Levels) trong SQL và các lỗi đọc dữ liệu tương ứng (Dirty Read, Non-repeatable Read, Phantom Read).",
            "Sự khác biệt giữa WHERE và HAVING trong SQL là gì? Cho ví dụ cụ thể khi nào bắt buộc phải dùng HAVING.",
            "Giải thích cơ chế hoạt động của INNER JOIN, LEFT JOIN, RIGHT JOIN và FULL OUTER JOIN. Làm thế nào để xử lý các giá trị NULL khi sử dụng JOIN?",
            "Trình bày về kỹ thuật Phân vùng bảng (Table Partitioning) và Phân mảnh dữ liệu (Sharding). Khi nào nên áp dụng từng kỹ thuật này?"
        ],
        "javascript": [
            "Closure trong JavaScript là gì, và chúng giúp đóng gói dữ liệu như thế nào?",
            "Giải thích cơ chế Event Loop trong JavaScript. Khác biệt giữa Microtask Queue và Macrotask Queue là gì?",
            "Sự khác biệt giữa var, let và const là gì? Giải thích khái niệm Hoisting và Temporal Dead Zone trong ES6.",
            "Trình bày sự khác biệt giữa Promise.all, Promise.allSettled, Promise.any và Promise.race. Cho ví dụ thực tế khi sử dụng.",
            "Prototype trong JavaScript là gì? Cơ chế kế thừa prototype (prototypal inheritance) hoạt động như thế nào?"
        ],
        "react": [
            "Giải thích vòng đời (lifecycle) của component trong React và sự khác biệt giữa state và props.",
            "Virtual DOM trong React hoạt động như thế nào? Thuật toán so khớp (reconciliation/diffing algorithm) tối ưu hiệu năng hiển thị ra sao?",
            "Sự khác biệt giữa useMemo và useCallback là gì? Khi nào nên sử dụng chúng và khi nào việc lạm dụng chúng sẽ gây phản tác dụng?",
            "Giải thích các cách quản lý State toàn cục (Global State Management) trong React. So sánh giữa React Context API và Redux.",
            "React Server Components (RSC) là gì? Khác biệt cơ bản giữa RSC và Client Components về mặt hiệu năng và cách render?"
        ],
        "machine learning": [
            "Làm thế nào để giải quyết vấn đề quá khớp (overfitting) trong các mô hình học máy?",
            "Giải thích sự khác biệt giữa các thuật toán học có giám sát (Supervised Learning) và học không giám sát (Unsupervised Learning). Cho ví dụ thực tế.",
            "Trình bày ý nghĩa của các độ đo Precision, Recall, F1-Score và ROC-AUC. Khi nào bạn ưu tiên Recall hơn Precision?",
            "Gradient Descent là gì? Giải thích sự khác biệt giữa Batch Gradient Descent, Stochastic Gradient Descent (SGD) và Mini-batch Gradient Descent.",
            "Làm thế nào để xử lý tập dữ liệu bị mất cân bằng lớp (imbalanced dataset) khi huấn luyện mô hình phân loại?"
        ],
        "fastapi": [
            "FastAPI xử lý các yêu cầu bất đồng bộ như thế nào, và ưu điểm hiệu năng của nó là gì?",
            "Dependency Injection (DI) trong FastAPI hoạt động như thế nào? Hãy nêu một trường hợp thực tế bạn dùng DI để quản lý tài nguyên hoặc bảo mật.",
            "Làm thế nào để tích hợp cơ chế xác thực JWT (JSON Web Token) và phân quyền người dùng trong ứng dụng FastAPI?",
            "Trình bày cách xử lý lỗi tập trung (Global Exception Handling) và tối ưu hóa việc validate dữ liệu bằng Pydantic trong FastAPI.",
            "FastAPI hỗ trợ Background Tasks như thế nào? Khi nào bạn sử dụng BackgroundTasks của FastAPI và khi nào nên dùng Celery?"
        ],
        "git": [
            "Sự khác biệt giữa git merge và git rebase là gì, và khi nào bạn nên sử dụng từng loại?",
            "Giải thích cơ chế hoạt động của Gitflow và Trunk-based Development. Dự án của bạn đang áp dụng mô hình nào và tại sao?",
            "Làm thế nào để giải quyết xung đột mã nguồn (Merge Conflict) trong Git một cách an toàn? Giải thích lệnh git stash dùng trong trường hợp nào.",
            "Sự khác biệt giữa git reset --soft, --mixed và --hard là gì? Bạn xử lý như thế nào nếu lỡ tay commit nhầm mã nguồn nhạy cảm (như API Key) lên remote repository?",
            "Lệnh git cherry-pick được dùng khi nào? Hãy mô tả một kịch bản thực tế cần dùng đến lệnh này."
        ],
        "docker": [
            "Docker Image và Docker Container khác nhau như thế nào? Trình bày cơ chế Union File System (UnionFS) trong Docker.",
            "Làm thế nào để tối ưu hóa kích thước của Docker Image (ví dụ: sử dụng Multi-stage build, Alpine base image)?",
            "Sự khác biệt giữa Docker Volumes và Bind Mounts là gì? Khi nào bạn sử dụng từng loại để lưu trữ dữ liệu bền vững?",
            "Giải thích các chế độ mạng (Network Drivers) trong Docker (Bridge, Host, Overlay, None). Khi nào nên dùng Overlay Network?",
            "Docker Compose là gì? Làm thế nào để quản lý các biến môi trường và thiết lập thứ tự khởi động dịch vụ bằng depends_on?"
        ],
        "kubernetes": [
            "Pod trong Kubernetes là gì? Tại sao k8s sử dụng Pod thay vì chạy trực tiếp Container?",
            "Sự khác biệt giữa Deployment và StatefulSet trong Kubernetes là gì? Khi nào bắt buộc phải dùng StatefulSet?",
            "Giải thích cách hoạt động của Kubernetes Service và các loại Service: ClusterIP, NodePort, LoadBalancer, ExternalName.",
            "Ingress Controller trong Kubernetes là gì và nó khác biệt gì so với một LoadBalancer Service thông thường?",
            "Cơ chế tự động co giãn (HPA - Horizontal Pod Autoscaler) hoạt động như thế nào? Độ đo nào thường được cấu hình để kích hoạt HPA?"
        ],
        "k8s": [
            "Pod trong Kubernetes là gì? Tại sao k8s sử dụng Pod thay vì chạy trực tiếp Container?",
            "Sự khác biệt giữa Deployment và StatefulSet trong Kubernetes là gì? Khi nào bắt buộc phải dùng StatefulSet?",
            "Giải thích cách hoạt động của Kubernetes Service và các loại Service: ClusterIP, NodePort, LoadBalancer, ExternalName.",
            "Ingress Controller trong Kubernetes là gì và nó khác biệt gì so với một LoadBalancer Service thông thường?",
            "Cơ chế tự động co giãn (HPA - Horizontal Pod Autoscaler) hoạt động như thế nào? Độ đo nào thường được cấu hình để kích hoạt HPA?"
        ],
        "aws": [
            "Giải thích sự khác biệt giữa AWS IAM Role và IAM User. Khi nào bạn nên gán IAM Role cho EC2 instance thay vì dùng IAM Access Key?",
            "Trình bày cơ chế hoạt động của Amazon S3 và các storage classes phổ biến (Standard, Infrequent Access, Glacier). Làm thế nào để tối ưu chi phí lưu trữ trên S3?",
            "Thiết kế một hệ thống mạng VPC (Virtual Private Cloud) trên AWS đảm bảo tính bảo mật cho ứng dụng web (gồm Public Subnet, Private Subnet, NAT Gateway).",
            "AWS Lambda hoạt động theo cơ chế Serverless như thế nào? Làm thế nào để hạn chế vấn đề khởi động nguội (Cold Start) trong Lambda?",
            "Khác biệt giữa Amazon RDS (Relational Database) và DynamoDB (NoSQL) về mặt mở rộng, cấu trúc dữ liệu và trường hợp sử dụng tối ưu."
        ],
        "amazon web services": [
            "Giải thích sự khác biệt giữa AWS IAM Role và IAM User. Khi nào bạn nên gán IAM Role cho EC2 instance thay vì dùng IAM Access Key?",
            "Trình bày cơ chế hoạt động của Amazon S3 và các storage classes phổ biến (Standard, Infrequent Access, Glacier). Làm thế nào để tối ưu chi phí lưu trữ trên S3?",
            "Thiết kế một hệ thống mạng VPC (Virtual Private Cloud) trên AWS đảm bảo tính bảo mật cho ứng dụng web (gồm Public Subnet, Private Subnet, NAT Gateway).",
            "AWS Lambda hoạt động theo cơ chế Serverless như thế nào? Làm thế nào để hạn chế vấn đề khởi động nguội (Cold Start) trong Lambda?",
            "Khác biệt giữa Amazon RDS (Relational Database) và DynamoDB (NoSQL) về mặt mở rộng, cấu trúc dữ liệu và trường hợp sử dụng tối ưu."
        ],
        "azure": [
            "Azure Active Directory (Azure AD) là gì và nó hỗ trợ quản lý danh tính, phân quyền (RBAC) cho tài nguyên Azure như thế nào?",
            "So sánh sự khác biệt giữa Azure Virtual Machines, App Services và Azure Kubernetes Service (AKS) về mức độ quản trị và khả năng mở rộng.",
            "Azure Blob Storage là gì? Giải thích các loại Blob (Block, Page, Append) và các mức truy cập (Hot, Cool, Cold, Archive).",
            "Làm thế nào để thiết lập kết nối an toàn giữa các dịch vụ trong Azure bằng Azure Private Link / Private Endpoint?",
            "Azure Functions là gì? Trình bày cơ chế kích hoạt (Triggers) và kết nối (Bindings) giúp lập trình serverless trên Azure hiệu quả."
        ],
        "microsoft azure": [
            "Azure Active Directory (Azure AD) là gì và nó hỗ trợ quản lý danh tính, phân quyền (RBAC) cho tài nguyên Azure như thế nào?",
            "So sánh sự khác biệt giữa Azure Virtual Machines, App Services và Azure Kubernetes Service (AKS) về mức độ quản trị và khả năng mở rộng.",
            "Azure Blob Storage là gì? Giải thích các loại Blob (Block, Page, Append) và các mức truy cập (Hot, Cool, Cold, Archive).",
            "Làm thế nào để thiết lập kết nối an toàn giữa các dịch vụ trong Azure bằng Azure Private Link / Private Endpoint?",
            "Azure Functions là gì? Trình bày cơ chế kích hoạt (Triggers) và kết nối (Bindings) giúp lập trình serverless trên Azure hiệu quả."
        ],
        "devops": [
            "Mục tiêu cốt lõi của văn hóa DevOps là gì? Hãy trình bày các chỉ số đo lường hiệu quả DevOps (DORA metrics) như Lead Time, Deployment Frequency, MTTR, CFR.",
            "Giải thích khái niệm Hạ tầng dưới dạng mã (IaC - Infrastructure as Code). So sánh Declarative (như Terraform) và Imperative (như Ansible).",
            "Làm thế nào để xây dựng một quy trình giám sát (Monitoring & Logging) hiệu quả cho ứng dụng Microservices? Bạn sử dụng các công cụ nào (như Prometheus, Grafana, ELK Stack)?",
            "Giải thích cơ chế Zero-downtime Deployment. So sánh các chiến lược triển khai: Blue-Green Deployment và Canary Deployment.",
            "Mô tả cách bạn tích hợp bảo mật vào quy trình DevOps (DevSecOps) ngay từ giai đoạn viết code đến triển khai thực tế."
        ],
        "ci/cd": [
            "Sự khác biệt giữa Tích hợp liên tục (Continuous Integration) và Triển khai liên tục (Continuous Delivery/Deployment) là gì?",
            "Hãy thiết kế một pipeline CI/CD hoàn chỉnh cho một ứng dụng web (bao gồm Linting, Unit Test, Build Image, Scan Security, Deploy) bằng GitHub Actions hoặc GitLab CI.",
            "Làm thế nào để quản lý các thông tin nhạy cảm (Secrets/Credentials) như DB password, API keys một cách an toàn trong CI/CD pipeline?",
            "Trình bày cơ chế Cache trong CI/CD pipeline giúp tăng tốc độ build và test ứng dụng.",
            "GitOps là gì? Nó cải tiến quy trình CD truyền thống như thế nào (ví dụ sử dụng ArgoCD hoặc FluxCD)?"
        ],
        "continuous integration": [
            "Sự khác biệt giữa Tích hợp liên tục (Continuous Integration) và Triển khai liên tục (Continuous Delivery/Deployment) là gì?",
            "Hãy thiết kế một pipeline CI/CD hoàn chỉnh cho một ứng dụng web (bao gồm Linting, Unit Test, Build Image, Scan Security, Deploy) bằng GitHub Actions hoặc GitLab CI.",
            "Làm thế nào để quản lý các thông tin nhạy cảm (Secrets/Credentials) như DB password, API keys một cách an toàn trong CI/CD pipeline?",
            "Trình bày cơ chế Cache trong CI/CD pipeline giúp tăng tốc độ build và test ứng dụng.",
            "GitOps là gì? Nó cải tiến quy trình CD truyền thống như thế nào (ví dụ sử dụng ArgoCD hoặc FluxCD)?"
        ],
        "spark": [
            "Apache Spark là gì và tại sao nó lại xử lý dữ liệu nhanh hơn MapReduce của Hadoop (ví dụ: In-memory processing, Lazy evaluation)?",
            "Giải thích khái niệm RDD (Resilient Distributed Dataset), DataFrame và Dataset trong Spark. Khác biệt chính giữa chúng là gì?",
            "Sự khác biệt giữa Transformation và Action trong Spark là gì? Cho ví dụ và giải thích cơ chế Directed Acyclic Graph (DAG).",
            "Trình bày vấn đề lệch dữ liệu (Data Skew) trong Spark. Làm thế nào để phát hiện và xử lý hiện tượng này để tránh nghẽn luồng (stragglers)?",
            "Khác biệt giữa Narrow Dependency và Wide Dependency trong Spark là gì? Tại sao Wide Dependency lại gây tốn tài nguyên mạng (Shuffle)?"
        ],
        "apache spark": [
            "Apache Spark là gì và tại sao nó lại xử lý dữ liệu nhanh hơn MapReduce của Hadoop (ví dụ: In-memory processing, Lazy evaluation)?",
            "Giải thích khái niệm RDD (Resilient Distributed Dataset), DataFrame và Dataset trong Spark. Khác biệt chính giữa chúng là gì?",
            "Sự khác biệt giữa Transformation và Action trong Spark là gì? Cho ví dụ và giải thích cơ chế Directed Acyclic Graph (DAG).",
            "Trình bày vấn đề lệch dữ liệu (Data Skew) trong Spark. Làm thế nào để phát hiện và xử lý hiện tượng này để tránh nghẽn luồng (stragglers)?",
            "Khác biệt giữa Narrow Dependency và Wide Dependency trong Spark là gì? Tại sao Wide Dependency lại gây tốn tài nguyên mạng (Shuffle)?"
        ],
        "nosql": [
            "Khác biệt cơ bản giữa cơ sở dữ liệu quan hệ (SQL) và phi quan hệ (NoSQL). Khi nào bạn nên chọn NoSQL thay vì SQL?",
            "Giải thích định lý CAP (CAP Theorem) trong hệ thống phân tán. Cơ sở dữ liệu NoSQL bạn chọn (như MongoDB hoặc Cassandra) ưu tiên cặp thuộc tính nào?",
            "Trình bày cách thiết kế dữ liệu trong MongoDB: khi nào nên nhúng dữ liệu (Embedding) và khi nào nên tham chiếu dữ liệu (Referencing/Normalization)?",
            "Trình bày cơ chế phân tán dữ liệu của Apache Cassandra. Làm thế nào để đảm bảo tính nhất quán dữ liệu (Tunable Consistency)?",
            "Redis là gì? Trình bày các kiểu dữ liệu nâng cao trong Redis và các cơ chế ghi dữ liệu xuống đĩa (AOF vs. RDB)."
        ],
        "system design": [
            "Trình bày các chiến lược thiết kế bộ nhớ đệm (Caching strategies) như Cache-Aside, Write-Through, Write-Behind. Khi nào sử dụng từng loại?",
            "Làm thế nào để thiết kế một hệ thống chịu tải cao và có khả năng mở rộng ngang (Horizontal Scaling)? Trình bày vai trò của Load Balancer và Database Replication.",
            "Giải thích cách hoạt động của cơ chế API Gateway và Rate Limiting. Làm thế nào để phòng chống tấn công DDoS ở tầng ứng dụng?",
            "Thiết kế một hệ thống hàng đợi thông điệp (Message Queue) sử dụng Kafka hoặc RabbitMQ. Sự khác biệt giữa mô hình Pub/Sub và Message Queue truyền thống là gì?",
            "Làm thế nào để đảm bảo tính nhất quán dữ liệu trong kiến trúc Microservices khi một giao dịch trải dài qua nhiều service (Saga Pattern)?"
        ],
        "c++": [
            "Sự khác biệt giữa con trỏ thường (raw pointer) và con trỏ thông minh (smart pointer - std::unique_ptr, std::shared_ptr, std::weak_ptr) trong C++11 là gì?",
            "Giải thích cơ chế quản lý bộ nhớ của C++ bao gồm Stack và Heap. Hiện tượng rò rỉ bộ nhớ (memory leak) xảy ra khi nào và cách khắc phục?",
            "Đa hình (Polymorphism) trong C++ được thực hiện như thế nào? Giải thích cơ chế bảng ảo (vtable) và con trỏ bảng ảo (vptr).",
            "Sự khác biệt giữa nạp chồng phương thức (Function Overloading) và ghi đè phương thức (Function Overriding) trong C++.",
            "Giải thích cơ chế Rvalue Reference và Move Semantics trong C++11. Tại sao chúng giúp tối ưu hiệu năng đáng kể?"
        ],
        "c#": [
            "Cơ chế quản lý bộ nhớ và bộ thu gom rác (Garbage Collector - GC) trong .NET/C# hoạt động như thế nào? Sự khác biệt giữa các thế hệ GC (Generation 0, 1, 2) là gì?",
            "Sự khác biệt giữa struct (Value Type) và class (Reference Type) trong C# là gì? Giải thích khái niệm Boxing và Unboxing.",
            "Khác biệt giữa từ khóa async / await và việc sử dụng lớp Thread truyền thống trong C# là gì?",
            "LINQ (Language Integrated Query) trong C# là gì? Khác biệt giữa IEnumerable và IQueryable khi truy vấn dữ liệu từ DB.",
            "Giải thích cơ chế ủy quyền (Delegates) và sự kiện (Events) trong C#. Khác biệt giữa Action, Func và Predicate."
        ],
        "typescript": [
            "TypeScript giải quyết những hạn chế nào của JavaScript? Sự khác biệt giữa interface và type alias trong TypeScript là gì?",
            "Giải thích các khái niệm nâng cao trong TypeScript: Generics, Utility Types (như Partial, Readonly, Pick, Omit) và Union/Intersection types.",
            "Khái niệm 'Type Assertion' và 'Type Guard' trong TypeScript là gì? Cho ví dụ về cách tự viết một custom type guard function.",
            "Chế độ 'strict' trong tsconfig.json mang lại lợi ích gì? Nêu ý nghĩa của cấu hình noImplicitAny và strictNullChecks.",
            "Decorator trong TypeScript hoạt động ra sao và được ứng dụng thế như thế nào trong các framework như Angular hoặc NestJS?"
        ],
        "html & css": [
            "Giải thích mô hình hộp (Box Model) trong CSS. Sự khác biệt giữa box-sizing: content-box' và box-sizing: border-box' là gì?",
            "Trình bày sự khác biệt giữa Flexbox và CSS Grid. Khi nào bạn ưu tiên sử dụng từng loại để thiết kế bố cục giao diện?",
            "Làm thế nào để tối ưu hóa hiệu năng tải trang web (Web Performance) liên quan đến HTML/CSS (ví dụ: Critical CSS, lazy load, tối ưu font, minify)?",
            "Giải thích cơ chế xếp chồng (Z-index và Stacking Context) trong CSS. Tại sao đôi khi đặt z-index rất cao nhưng phần tử vẫn bị che khuất?",
            "Ý nghĩa của Responsive Web Design là gì? Khác biệt giữa Media Queries và việc sử dụng các đơn vị linh hoạt như 'em', 'rem', 'vh', 'vw'."
        ],
        "html": [
            "Giải thích mô hình hộp (Box Model) trong CSS. Sự khác biệt giữa box-sizing: content-box' và box-sizing: border-box' là gì?",
            "Trình bày sự khác biệt giữa Flexbox và CSS Grid. Khi nào bạn ưu tiên sử dụng từng loại để thiết kế bố cục giao diện?",
            "Làm thế nào để tối ưu hóa hiệu năng tải trang web (Web Performance) liên quan đến HTML/CSS (ví dụ: Critical CSS, lazy load, tối ưu font, minify)?",
            "Giải thích cơ chế xếp chồng (Z-index và Stacking Context) trong CSS. Tại sao đôi khi đặt z-index rất cao nhưng phần tử vẫn bị che khuất?",
            "Ý nghĩa của Responsive Web Design là gì? Khác biệt giữa Media Queries và việc sử dụng các đơn vị linh hoạt như 'em', 'rem', 'vh', 'vw'."
        ],
        "css": [
            "Giải thích mô hình hộp (Box Model) trong CSS. Sự khác biệt giữa box-sizing: content-box' và box-sizing: border-box' là gì?",
            "Trình bày sự khác biệt giữa Flexbox và CSS Grid. Khi nào bạn ưu tiên sử dụng từng loại để thiết kế bố cục giao diện?",
            "Làm thế nào để tối ưu hóa hiệu năng tải trang web (Web Performance) liên quan đến HTML/CSS (ví dụ: Critical CSS, lazy load, tối ưu font, minify)?",
            "Giải thích cơ chế xếp chồng (Z-index và Stacking Context) trong CSS. Tại sao đôi khi đặt z-index rất cao nhưng phần tử vẫn bị che khuất?",
            "Ý nghĩa của Responsive Web Design là gì? Khác biệt giữa Media Queries và việc sử dụng các đơn vị linh hoạt như 'em', 'rem', 'vh', 'vw'."
        ],
        "go": [
            "Goroutine trong Go là gì và tại sao nó lại nhẹ hơn rất nhiều so với Thread của hệ điều hành?",
            "Giải thích cơ chế giao tiếp giữa các Goroutine bằng Channel. Khác biệt giữa Buffered Channel và Unbuffered Channel là gì?",
            "Cơ chế quản lý lỗi (Error Handling) trong Go có gì đặc biệt so với các ngôn ngữ có cơ chế try-catch? Tại sao Go lại thiết kế như vậy?",
            "Trình bày cơ chế hoạt động của từ khóa 'defer'. Thứ tự thực thi của nhiều lệnh 'defer' trong một hàm là gì?",
            "Sự khác biệt giữa con trỏ (pointer) trong Go và con trỏ trong C/C++ là gì? Go có cho phép tính toán con trỏ (pointer arithmetic) không?"
        ],
        "golang": [
            "Goroutine trong Go là gì và tại sao nó lại nhẹ hơn rất nhiều so với Thread của hệ điều hành?",
            "Giải thích cơ chế giao tiếp giữa các Goroutine bằng Channel. Khác biệt giữa Buffered Channel và Unbuffered Channel là gì?",
            "Cơ chế quản lý lỗi (Error Handling) trong Go có gì đặc biệt so với các ngôn ngữ có cơ chế try-catch? Tại sao Go lại thiết kế như vậy?",
            "Trình bày cơ chế hoạt động của từ khóa 'defer'. Thứ tự thực thi của nhiều lệnh 'defer' trong một hàm là gì?",
            "Sự khác biệt giữa con trỏ (pointer) trong Go và con trỏ trong C/C++ là gì? Go có cho phép tính toán con trỏ (pointer arithmetic) không?"
        ],
        "nodejs": [
            "Node.js là đơn luồng (single-threaded) hay đa luồng (multi-threaded)? Giải thích cơ chế Non-blocking I/O và Event Loop trong Node.js.",
            "Sự khác biệt giữa 'require' (CommonJS) và 'import' (ES Modules) trong Node.js là gì?",
            "Trình bày cách quản lý các luồng dữ liệu lớn bằng Stream trong Node.js. Nêu ví dụ về việc đọc một tệp dung lượng cực lớn mà không bị tràn bộ nhớ (Out of Memory).",
            "Middleware trong Express.js là gì? Hãy nêu cơ chế hoạt động của chuỗi middleware và vai trò của hàm 'next()'.",
            "Làm thế nào để xử lý các Unhandled Rejections và Uncaught Exceptions trong Node.js để tránh sập server đột ngột?"
        ],
        "node.js": [
            "Node.js là đơn luồng (single-threaded) hay đa luồng (multi-threaded)? Giải thích cơ chế Non-blocking I/O và Event Loop trong Node.js.",
            "Sự khác biệt giữa 'require' (CommonJS) và 'import' (ES Modules) trong Node.js là gì?",
            "Trình bày cách quản lý các luồng dữ liệu lớn bằng Stream trong Node.js. Nêu ví dụ về việc đọc một tệp dung lượng cực lớn mà không bị tràn bộ nhớ (Out of Memory).",
            "Middleware trong Express.js là gì? Hãy nêu cơ chế hoạt động của chuỗi middleware và vai trò của hàm 'next()'.",
            "Làm thế nào để xử lý các Unhandled Rejections và Uncaught Exceptions trong Node.js để tránh sập server đột ngột?"
        ]
    }
    
    # Shuffle user skills to randomize matching order
    skills_shuffled = list(skills)
    random.shuffle(skills_shuffled)
    
    matched_questions = []
    for s in skills_shuffled:
        s_low = s.lower()
        if s_low in skills_questions:
            selected_q = random.choice(skills_questions[s_low])
            if selected_q not in matched_questions:
                matched_questions.append(selected_q)
            if len(matched_questions) >= 3: # Up to 3 technical questions based on skills
                break
                
    # Standard fallback behavioral/situational questions (translated to Vietnamese)
    fallback_questions = [
        "Bạn có thể mô tả một dự án đầy thử thách mà bạn đã tham gia gần đây không? Vai trò của bạn là gì và kết quả ra sao?",
        "Hãy kể về một tình huống bạn có bất đồng quan điểm với thành viên trong nhóm hoặc đối tác. Bạn đã giải quyết nó như thế nào?",
        "Làm thế nào để bạn sắp xếp thứ tự ưu tiên cho các công việc khi làm việc trên nhiều dự án có thời hạn gấp rút?",
        "Hãy kể lại một lần bạn phải học một công nghệ hoặc công cụ mới một cách nhanh chóng. Bạn đã tiếp cận quá trình học tập đó như thế nào?",
        "Tại sao bạn muốn gia nhập công ty chúng tôi, và vai trò này phù hợp như thế nào với mục tiêu nghề nghiệp dài hạn của bạn?",
        "Kể về một lần bạn mắc lỗi trong công việc (ví dụ: làm chậm tiến độ, hoặc deploy lỗi). Bạn đã xử lý hậu quả thế nào và rút ra bài học gì?",
        "Khi nhận được những ý kiến đóng góp hoặc phê bình tiêu cực từ đồng nghiệp hoặc cấp trên, bạn thường tiếp nhận và phản hồi như thế nào?",
        "Mô tả một tình huống bạn phải làm việc với một khách hàng hoặc thành viên nhóm rất khó tính. Bạn đã làm gì để duy trì tiến độ công việc?",
        "Hãy chia sẻ về một quyết định kỹ thuật/thiết kế quan trọng mà bạn đã đưa ra trong dự án cũ. Bạn đã thuyết phục mọi người như thế nào?",
        "Làm thế nào để bạn tự tạo động lực và duy trì năng suất làm việc khi phải thực hiện các tác vụ lặp đi lặp lại hoặc nhàm chán?",
        "Khi một dự án bị trễ hạn do các yếu tố khách quan ngoài tầm kiểm soát của bạn, bạn sẽ xử lý tình huống đó thế nào với các bên liên quan?",
        "Hãy mô tả cách bạn hướng dẫn hoặc hỗ trợ (mentor) một đồng nghiệp mới hoặc một bạn cấp dưới (junior) hòa nhập và làm việc hiệu quả.",
        "Kể về một dự án mà mục tiêu ban đầu liên tục bị thay đổi (scope creep). Bạn đã làm thế nào để thích ứng và đưa dự án về đích?",
        "Theo bạn, điều gì là quan trọng nhất để xây dựng và duy trì một văn hóa làm việc nhóm (teamwork) tích cực, gắn kết?",
        "Hãy kể về một sáng kiến hoặc cải tiến mà bạn tự đề xuất giúp nâng cao hiệu suất làm việc hoặc giảm thiểu lỗi trong dự án.",
        "Khi phải đưa ra quyết định mà không có đầy đủ thông tin hoặc dữ liệu cần thiết, bạn sẽ tiếp cận và đánh giá rủi ro như thế nào?"
    ]
    
    needed = 5 - len(matched_questions)
    sampled_fallbacks = random.sample(fallback_questions, min(needed, len(fallback_questions)))
    final_questions = matched_questions + sampled_fallbacks
    
    # Pad if not enough questions (just to be completely safe)
    if len(final_questions) < 5:
        for q in fallback_questions:
            if q not in final_questions:
                final_questions.append(q)
            if len(final_questions) >= 5:
                break
                
    # Shuffle final 5 questions to mix technical and behavioral questions
    random.shuffle(final_questions)
    return json.dumps(final_questions[:5], indent=2, ensure_ascii=False)
 
def check_answer_relevance(question, answer, skills=None):
    # Clean inputs
    q_clean = question.lower()
    a_clean = answer.lower()
    
    # Core stopwords for Vietnamese questions
    stopwords = {
        "bạn", "có", "thể", "một", "mà", "đã", "gần", "đây", "không", "của", "là", 
        "gì", "và", "ra", "sao", "hãy", "kể", "về", "tình", "huống", "bất", "đồng", 
        "quan", "điểm", "với", "thành", "viên", "trong", "nhóm", "đối", "tác", "như", 
        "thế", "nào", "ở", "để", "cho", "này", "câu", "hỏi", "được", "làm", "tại", "sự", "khi",
        "nhiệm", "vụ", "vai", "trò", "kết", "quả", "mô", "tả", "những", "các", "như", "thế", "nào",
        "tôi", "chúng", "ta", "từng", "lần", "phải", "cách", "hạn", "đối", "tác", "thành", "viên",
        "trong", "nhóm", "với", "trên", "nhiều", "gấp", "rút", "mới", "nhanh", "chóng", "muốn", 
        "phù", "hợp", "dài", "trước",
        "gia", "khác", "biệt", "giữa", "đầy", "thử", "thách", "tham", "giải", "quyết", 
        "nó", "sắp", "xếp", "thứ", "tự", "ưu", "tiên", "học", "tiếp", "cận", "quá", 
        "trình", "nhập", "mục", "tiêu", "nghề", "nghiệp", "công", "việc", "nên", "sử", "dụng", "dùng"
    }
    
    # Tokenize and clean question
    q_words = [w.strip("?,.:;!\"()[]{}") for w in q_clean.split()]
    q_content_words = [w for w in q_words if w and w not in stopwords]
    
    # Tokenize and clean answer
    a_words = [w.strip("?,.:;!\"()[]{}") for w in a_clean.split()]
    a_content_words = [w for w in a_words if w and w not in stopwords]
    
    if not q_content_words or not a_content_words:
        return False
        
    # Check direct word overlap
    overlap = set(q_content_words).intersection(set(a_content_words))
    if len(overlap) >= 1:
        return True
        
    # Common tech keywords
    tech_keywords = {
        "python", "sql", "git", "docker", "kubernetes", "aws", "azure", "gcp", 
        "nosql", "hadoop", "spark", "mapreduce", "etl", "api", "django", "flask", 
        "react", "java", "c++", "c#", "javascript", "typescript", "html", "css", 
        "linux", "bash", "shell", "list", "tuple", "generator", "airflow", "kafka",
        "redshift", "bigquery", "postgres", "mysql", "mongodb", "cassandra"
    }
    
    q_tech = [w for w in q_content_words if w in tech_keywords]
    a_tech = [w for w in a_content_words if w in tech_keywords]
    
    if q_tech:
        # Technical question: answer must contain at least one related tech keyword
        overlap_tech = set(q_tech).intersection(set(a_tech))
        if len(overlap_tech) >= 1:
            return True
        return False
        
    # Behavioral question: check if the answer discusses professional/workplace terms
    work_keywords = {
        "dự án", "công việc", "khách hàng", "đồng nghiệp", "quản lý", "phát triển", 
        "thiết kế", "xử lý", "hệ thống", "tối ưu", "giải quyết", "khó khăn", "thử thách", 
        "thành viên", "nhóm", "leader", "sản phẩm", "yêu cầu", "tiến độ", "kỹ năng", 
        "học tập", "công nghệ", "công cụ", "đối tác", "quy trình", "vận hành"
    }
    
    # Check if any work keyword is in the answer
    for kw in work_keywords:
        if kw in a_clean:
            return True
            
    # Check skills exact matching
    if skills:
        a_words_set = set(a_words)
        for s in skills:
            s_words = s.lower().split()
            if len(s_words) == 1:
                if s_words[0] in a_words_set:
                    return True
            else:
                if f" {s.lower()} " in f" {a_clean} ":
                    return True
                    
    return False

def evaluate_local_answers(resume_text, jd_text, questions, answers):
    """Local evaluation of mock interview answers based on word count and keyword matching in Vietnamese."""
    global nlp
    if nlp is None:
        nlp = load_spacy_nlp_model()
        
    doc = nlp(resume_text)
    skills = [s.lower() for s in extract_skills_fuzzy(doc)]
    
    evaluations = []
    total_score = 0
    
    for q, a in zip(questions, answers):
        a_clean = (a or "").strip().lower()
        words = a_clean.split()
        word_count = len(words)
        
        # Heuristic scoring based on length, relevance and content
        is_relevant = check_answer_relevance(q, a_clean, skills)
        
        if word_count == 0:
            score = 0
            feedback = "Không có câu trả lời nào được cung cấp. Trong một cuộc phỏng vấn thực tế, bạn luôn nên cố gắng trả lời, ngay cả khi chỉ để mô tả quá trình suy nghĩ của mình."
            excellent = "Cung cấp một câu trả lời có cấu trúc bằng phương pháp STAR: Tình huống (Situation), Nhiệm vụ (Task), Hành động (Action), Kết quả (Result). Ví dụ: 'Trong vai trò trước đây của tôi, chúng tôi đã đối mặt với [Tình huống]... Nhiệm vụ của tôi là [Nhiệm vụ]... Tôi đã thực hiện hành động bằng cách [Hành động]... và kết quả là [Kết quả].'"
        elif not is_relevant:
            score = 1
            feedback = "Câu trả lời của bạn dường như lạc đề hoặc không đúng trọng tâm câu hỏi. Trong phỏng vấn thực tế, việc trả lời lạc đề sẽ bị đánh giá rất thấp. Hãy tập trung trả lời trực tiếp câu hỏi được hỏi."
            excellent = f"Để cấu trúc một câu trả lời xuất sắc cho câu hỏi này, hãy làm theo hướng dẫn sau:\n1. **Tình huống**: Mô tả một bối cảnh cụ thể mà câu hỏi này áp dụng.\n2. **Hành động**: Đề cập đến việc sử dụng các kỹ năng như {', '.join(skills[:3]) if skills else 'giải quyết vấn đề một cách phân tích'} để xử lý vấn đề.\n3. **Kết quả**: Kết luận với một kết quả kinh doanh hoặc kỹ thuật tích cực, rõ ràng."
        elif word_count < 5:
            score = 1
            feedback = "Câu trả lời của bạn quá ngắn (dưới 5 từ). Trong phỏng vấn thực tế, câu trả lời cực ngắn như vậy không thể truyền tải đủ thông tin và sẽ bị đánh giá rất thấp. Hãy cố gắng trả lời chi tiết và rõ ràng hơn."
            excellent = f"Để cấu trúc một câu trả lời xuất sắc cho câu hỏi này, hãy làm theo hướng dẫn sau:\n1. **Tình huống**: Mô tả một bối cảnh cụ thể mà câu hỏi này áp dụng.\n2. **Hành động**: Đề cập đến việc sử dụng các kỹ năng như {', '.join(skills[:3]) if skills else 'giải quyết vấn đề một cách phân tích'} để xử lý vấn đề.\n3. **Kết quả**: Kết luận với một kết quả kinh doanh hoặc kỹ thuật tích cực, rõ ràng."
        elif word_count < 15:
            score = 3
            feedback = "Câu trả lời còn khá ngắn và sơ sài (dưới 15 từ). Hãy cố gắng mở rộng thêm, đưa ra bối cảnh (Situation), hành động cụ thể (Action) hoặc kết quả (Result) của bạn để chứng minh năng lực."
            excellent = f"Để cấu trúc một câu trả lời xuất sắc cho câu hỏi này, hãy làm theo hướng dẫn sau:\n1. **Tình huống**: Mô tả một bối cảnh cụ thể mà câu hỏi này áp dụng.\n2. **Hành động**: Đề cập đến việc sử dụng các kỹ năng như {', '.join(skills[:3]) if skills else 'giải quyết vấn đề một cách phân tích'} để xử lý vấn đề.\n3. **Kết quả**: Kết luận với một kết quả kinh doanh hoặc kỹ thuật tích cực, rõ ràng."
        else:
            score = 5  # Reasonable base score for decent length
            # Word count contribution
            if word_count > 100:
                score += 3
            elif word_count > 60:
                score += 2
            elif word_count > 30:
                score += 1
                
            # Keyword/Skill match contribution
            matched_skills_in_answer = [s for s in skills if s in a_clean]
            if len(matched_skills_in_answer) >= 2:
                score += 2
            elif len(matched_skills_in_answer) >= 1:
                score += 1
                
            score = min(score, 10)
            
            # Feedback generation (translated to Vietnamese)
            if word_count < 25:
                feedback = "Câu trả lời của bạn hơi ngắn. Hãy cố gắng nói chi tiết hơn về những đóng góp cụ thể của bạn, các công nghệ đã sử dụng và tác động từ hành động của bạn."
            else:
                feedback = "Nỗ lực tốt! Câu trả lời của bạn đã bao gồm các yếu tố cơ bản. Để cải thiện, hãy đảm bảo cấu trúc câu trả lời của bạn theo phương pháp STAR, làm nổi bật các kết quả định lượng (ví dụ: 'cải thiện hiệu suất thêm 20%') và liệt kê các công cụ cụ thể bạn đã sử dụng."
                
            excellent = f"Để cấu trúc một câu trả lời xuất sắc cho câu hỏi này, hãy làm theo hướng dẫn sau:\n1. **Tình huống**: Mô tả một bối cảnh cụ thể mà câu hỏi này áp dụng.\n2. **Hành động**: Đề cập đến việc sử dụng các kỹ năng như {', '.join(skills[:3]) if skills else 'giải quyết vấn đề một cách phân tích'} để xử lý vấn đề.\n3. **Kết quả**: Kết luận với một kết quả kinh doanh hoặc kỹ thuật tích cực, rõ ràng."
            
        total_score += score
        evaluations.append({
            "Question": q,
            "Answer": a or "",
            "Score": score,
            "Feedback": feedback,
            "ExcellentResponse": excellent
        })
        
    overall_score = round((total_score / (len(questions) * 10)) * 100) if questions else 0
    
    report = {
        "OverallScore": overall_score,
        "Evaluation": evaluations
    }
    return json.dumps(report, indent=2)

# --- PUBLIC INTERFACE ---

def generate_cover_letter_with_gemini(resume_text, jd_text, tone, api_key):
    if not api_key:
        return generate_local_cover_letter(resume_text, jd_text, tone)
        
    model_name = get_suitable_gemini_model(api_key)
    if not model_name:
        return generate_local_cover_letter(resume_text, jd_text, tone)
    
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(model_name)
        
        prompt = f"""
        You are an expert career consultant and resume writer. 
        Write a highly tailored and professional cover letter based on the candidate's resume and target Job Description (JD).
        Use a {tone} tone. Keep the letter structured, impactful, and under 300 words. Do not include placeholders (like [Date], [Company Name] if not found in the JD, or [Your Name] if found in the resume). Try to use actual details if available.
        IMPORTANT: The entire cover letter MUST be written in Vietnamese.
        
        Candidate Resume:
        {resume_text}
        
        Target Job Description:
        {jd_text}
        
        Generated Cover Letter (in Vietnamese):
        """
        
        response = model.generate_content(prompt)
        if response and response.text:
            return response.text.strip()
    except Exception:
        # Fallback to local on API error
        return generate_local_cover_letter(resume_text, jd_text, tone)
    return generate_local_cover_letter(resume_text, jd_text, tone)

def generate_interview_questions_with_gemini(resume_text, jd_text, api_key):
    if not api_key:
        return generate_local_interview_questions(resume_text, jd_text)
        
    model_name = get_suitable_gemini_model(api_key)
    if not model_name:
        return generate_local_interview_questions(resume_text, jd_text)
        
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(model_name)
        
        prompt = f"""
        You are an expert technical interviewer. Based on the candidate's resume and the job description (JD) below, generate 5 relevant, challenging, and specific interview questions (a mix of technical and situational/behavioral questions).
        IMPORTANT: The questions MUST be written in Vietnamese.
        
        Return a **single valid JSON array of strings**, where each string is a question in Vietnamese, and nothing else.
        
        Format:
        [
          "Question 1 in Vietnamese...",
          "Question 2 in Vietnamese...",
          "Question 3 in Vietnamese...",
          "Question 4 in Vietnamese...",
          "Question 5 in Vietnamese..."
        ]
        
        Rules:
        - Return ONLY the JSON array. Do not include markdown code block syntax (like ```json), commentary, or introductions.
        - Double quote all items.
        
        Candidate Resume:
        {resume_text}
        
        Target Job Description:
        {jd_text}
        """
        
        response = model.generate_content(prompt)
        if response and response.text:
            raw_text = response.text.strip()
            start = raw_text.find("[")
            end = raw_text.rfind("]") + 1
            if start != -1 and end != -1:
                raw_text = raw_text[start:end]
            return raw_text
    except Exception:
        return generate_local_interview_questions(resume_text, jd_text)
    return generate_local_interview_questions(resume_text, jd_text)

def evaluate_interview_answers_with_gemini(resume_text, jd_text, questions, answers, api_key):
    if not api_key:
        return evaluate_local_answers(resume_text, jd_text, questions, answers)
        
    model_name = get_suitable_gemini_model(api_key)
    if not model_name:
        return evaluate_local_answers(resume_text, jd_text, questions, answers)
        
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(model_name)
        
        q_a_payload = []
        for q, a in zip(questions, answers):
            q_a_payload.append({"question": q, "answer": a})
            
        prompt = f"""
        You are an expert mock interview assessor. Evaluate the user's answers to the interview questions based on their resume and target Job Description (JD).
        IMPORTANT: The evaluation feedback and Excellent Response MUST be written in Vietnamese.
        
        Provide constructive feedback, score (from 0 to 10), and a model "Excellent Response" for each question. All text descriptions ("Feedback" and "ExcellentResponse") must be in Vietnamese.
        Also, calculate an overall mock interview score from 0 to 100.
        
        Return a **single valid JSON object** and nothing else.
        
        Strict format:
        {{
          "OverallScore": 75,
          "Evaluation": [
            {{
              "Question": "Question text...",
              "Answer": "User's answer...",
              "Score": 8,
              "Feedback": "Feedback in Vietnamese...",
              "ExcellentResponse": "ExcellentResponse in Vietnamese using STAR method..."
            }},
            ...
          ]
        }}
        
        Rules:
        - Return ONLY the JSON object. Do not include markdown code block syntax (like ```json), preamble, or notes.
        - Ensure it is a valid parseable JSON.
        - CRITICAL RULES FOR SCORING AND FEEDBACK:
          * If the candidate's answer is empty, extremely short (e.g., under 5 words, single letters like 's', single words like 'không', 'ok', 'no'), or irrelevant/gibberish/off-topic (e.g. answering a Python coding question by talking about cooking, or talking about Docker when asked about Python lists/tuples), you MUST score it 0 or 1 out of 10. Do NOT give a high base score (like 5 or 6) for trivial, low-quality, or off-topic answers.
          * Give a score of 8 or above only to clear, detailed, professional answers that demonstrate relevant knowledge or experience and directly address the question.
          * The feedback and Excellent Response for each question must be written in Vietnamese.
        
        Candidate Resume:
        {resume_text}
        
        Target Job Description:
        {jd_text}
        
        Questions and Answers:
        {json.dumps(q_a_payload, indent=2)}
        """
        
        response = model.generate_content(prompt)
        if response and response.text:
            raw_text = response.text.strip()
            start = raw_text.find("{")
            end = raw_text.rfind("}") + 1
            if start != -1 and end != -1:
                raw_text = raw_text[start:end]
            return raw_text
    except Exception:
        return evaluate_local_answers(resume_text, jd_text, questions, answers)
    return evaluate_local_answers(resume_text, jd_text, questions, answers)

def generate_cover_letter_docx(name, email, phone, body_text):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO

    doc = Document()
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Letterhead style
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(name.upper())
    run.font.name = 'Arial'
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(13, 148, 136) # Teal
    
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c_run = contact_p.add_run(f"{email}   |   {phone}")
    c_run.font.name = 'Arial'
    c_run.font.size = Pt(10)
    c_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Add a divider line
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_line_run = p_line.add_run("__________________________________________________________________")
    p_line_run.font.color.rgb = RGBColor(220, 220, 220)
    
    doc.add_paragraph("") # Spacing
    
    # Add body paragraphs
    paragraphs = body_text.split("\n")
    for para in paragraphs:
        if para.strip():
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(para.strip())
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            
    # Save to BytesIO
    file_io = BytesIO()
    doc.save(file_io)
    file_io.seek(0)
    return file_io

SKILL_DETAILS_DB = {
    "Hadoop": {
        "task_title": "Lập trình & Vận hành Apache Hadoop",
        "sub_tasks": [
            "Cài đặt Hadoop Single-node Cluster cục bộ và làm quen với cấu hình core-site.xml, hdfs-site.xml.",
            "Thực hành sử dụng dòng lệnh HDFS CLI để quản lý tệp tin (hdfs dfs -put, -get, -ls, -rm).",
            "Viết và chạy một chương trình MapReduce đơn giản (WordCount) trên Yarn để hiểu cơ chế phân tán dữ liệu."
        ]
    },
    "Spark": {
        "task_title": "Xử lý Dữ liệu lớn với Apache Spark",
        "sub_tasks": [
            "Thiết lập môi trường làm việc với PySpark hoặc Spark Scala trên Jupyter Notebook.",
            "Thực hiện các thao tác chuyển đổi (map, filter, flatMap) và hành động (collect, count, reduce) trên RDD & DataFrames.",
            "Tối ưu hóa hiệu năng xử lý Spark SQL bằng cách sử dụng Partitioning, Bucketing và Broadcast Joins."
        ]
    },
    "ETL": {
        "task_title": "Thiết kế Luồng trích xuất & Chuyển đổi dữ liệu (ETL Pipeline)",
        "sub_tasks": [
            "Xây dựng pipeline trích xuất dữ liệu từ các API công cộng, chuyển đổi định dạng và làm sạch trường Null bằng Python.",
            "Thiết lập cơ sở dữ liệu đích (Data Warehouse) để nạp dữ liệu sạch vào cấu trúc bảng tối ưu.",
            "Lập lịch chạy tự động cho luồng ETL bằng Apache Airflow hoặc Cron Jobs, kèm theo cảnh báo email khi lỗi."
        ]
    },
    "AWS": {
        "task_title": "Thiết lập Hạ tầng Điện toán Đám mây AWS",
        "sub_tasks": [
            "Tạo tài khoản AWS Free Tier và cấu hình bảo mật IAM Users, Groups và Role với chính sách Principal least privilege.",
            "Khởi chạy EC2 Web Server Instance, cấu hình VPC, Subnets, Internet Gateways và Security Groups.",
            "Lưu trữ dữ liệu phi cấu trúc trên Amazon S3 và thiết lập cơ chế Lifecycle Rules để tối ưu chi phí."
        ]
    },
    "Azure": {
        "task_title": "Quản trị Dịch vụ Đám mây Microsoft Azure",
        "sub_tasks": [
            "Cấu hình Azure Virtual Machines và thiết lập Network Security Groups (NSG) để chặn lọc truy cập.",
            "Sử dụng Azure Blob Storage để quản lý tệp tin và Azure Resource Manager (ARM) để triển khai hạ tầng tự động.",
            "Cấu hình Azure Monitor để giám sát hiệu năng sử dụng CPU, RAM và băng thông của VM."
        ]
    },
    "NoSQL": {
        "task_title": "Quản trị Cơ sở Dữ liệu NoSQL (MongoDB, Cassandra)",
        "sub_tasks": [
            "Cài đặt và thiết lập MongoDB Community Server hoặc Apache Cassandra.",
            "Thiết kế schema dạng document-based, thực hành truy vấn phức tạp sử dụng Aggregation Framework.",
            "Tạo indexes (Single, Compound, Text) để cải thiện tốc độ truy vấn đối với các trường dữ liệu tìm kiếm nhiều."
        ]
    },
    "Big Data Architecture": {
        "task_title": "Kiến trúc Hệ thống Dữ liệu lớn",
        "sub_tasks": [
            "Nghiên cứu mô hình kiến trúc Kappa và Lambda để xử lý song song cả Batch Processing và Real-time Streaming.",
            "Vẽ sơ đồ luồng dữ liệu biểu diễn các tầng: Ingestion, Storage, Processing, Analytics và Serving.",
            "Tìm hiểu tiêu chuẩn tối ưu hóa thiết kế hệ thống dữ liệu có tính mở rộng cao (Scalability) và chịu lỗi tốt (Fault Tolerance)."
        ]
    },
    "Data Warehousing": {
        "task_title": "Xây dựng Kho Dữ liệu (Data Warehouse)",
        "sub_tasks": [
            "Thiết kế cấu trúc kho dữ liệu theo mô hình Star Schema hoặc Snowflake Schema với các bảng Fact và Dimension.",
            "Thực hành viết các truy vấn SQL phân tích phức tạp sử dụng các hàm Window Functions (ROW_NUMBER, Lead, Lag).",
            "Tìm hiểu cách tối ưu hóa hiệu năng truy vấn trên các kho dữ liệu lớn như Amazon Redshift hoặc Google BigQuery."
        ]
    },
    "Data Lakes": {
        "task_title": "Quản trị Hồ dữ liệu (Data Lakes)",
        "sub_tasks": [
            "Thiết lập Data Lake sử dụng AWS Lake Formation hoặc lưu trữ phân tán HDFS/S3.",
            "Phân loại cấu trúc thư mục lưu trữ theo phân vùng (Raw, Cleaned, Curated) để tránh biến hồ dữ liệu thành đầm lầy dữ liệu (Data Swamp).",
            "Cấu hình các công cụ cataloging dữ liệu như AWS Glue để tự động quét schema và lập chỉ mục metadata."
        ]
    },
    "Data Modeling": {
        "task_title": "Thiết kế & Mô hình hóa Dữ liệu",
        "sub_tasks": [
            "Vẽ mô hình quan hệ thực thể (ERD) cấp logic và vật lý bằng các công cụ chuyên dụng (dbdiagram.io, draw.io).",
            "Thực hiện chuẩn hóa dữ liệu (1NF, 2NF, 3NF) để giảm thiểu dư thừa dữ liệu và thiết kế phi chuẩn hóa khi cần tối ưu truy vấn.",
            "Xác định khóa chính (Primary Key), khóa ngoại (Foreign Key) và các ràng buộc dữ liệu để đảm bảo toàn vẹn dữ liệu."
        ]
    },
    "Data Governance": {
        "task_title": "Quản trị & Giám sát Chất lượng Dữ liệu (Data Governance)",
        "sub_tasks": [
            "Xây dựng chính sách bảo mật dữ liệu, phân quyền truy cập cột/hàng đối với dữ liệu nhạy cảm (PII).",
            "Thiết lập sơ đồ nguồn gốc dữ liệu (Data Lineage) để theo dõi vòng đời dữ liệu từ nguồn gốc đến báo cáo cuối cùng.",
            "Định nghĩa các chỉ số chất lượng dữ liệu (Data Quality Rules) và chạy kiểm tra định kỳ tính chính xác, nhất quán của dữ liệu."
        ]
    },
    "Cloud Computing": {
        "task_title": "Ứng dụng Công nghệ Điện toán Đám mây",
        "sub_tasks": [
            "Phân biệt và so sánh các mô hình dịch vụ đám mây chính: IaaS, PaaS, SaaS và mô hình triển khai Public, Private, Hybrid Cloud.",
            "Cài đặt và sử dụng Command Line Interface (CLI) của các nhà cung cấp đám mây để quản lý tài nguyên từ terminal.",
            "Tìm hiểu nguyên lý thanh toán theo mức độ sử dụng (Pay-as-you-go) và các chiến lược tối ưu chi phí hạ tầng."
        ]
    },
    "MapReduce": {
        "task_title": "Lập trình Tính toán Song song MapReduce",
        "sub_tasks": [
            "Hiểu sâu về cơ chế hoạt động của thuật toán MapReduce: Phase Map, Shuffle & Sort, Phase Reduce.",
            "Lập trình thuật toán lọc dữ liệu, tính giá trị trung bình hoặc kết nối dữ liệu (Join) phân tán sử dụng Java/Python.",
            "Tìm hiểu lý do các công nghệ hiện đại thay thế MapReduce bằng bộ nhớ trong (In-Memory) của Spark."
        ]
    },
    "Security": {
        "task_title": "Bảo mật & Mã hóa Hệ thống dữ liệu",
        "sub_tasks": [
            "Thiết lập giao thức truyền tải dữ liệu an toàn HTTPS, cấu hình chứng chỉ bảo mật SSL/TLS.",
            "Mã hóa dữ liệu tại chỗ (Encryption at rest) và dữ liệu đang truyền đi (Encryption in transit) bằng thuật toán AES/RSA.",
            "Cấu hình các bộ lọc tường lửa (Firewalls) và chính sách CORS để hạn chế truy cập trái phép vào tài nguyên API."
        ]
    },
    "Ansible": {
        "task_title": "Tự động hóa Hạ tầng với Ansible",
        "sub_tasks": [
            "Cài đặt Ansible trên máy Control Node, cấu hình tệp Inventory quản lý các Managed Nodes.",
            "Viết các Ansible Playbooks (.yml) để tự động hóa cài đặt phần mềm, phân quyền và khởi chạy dịch vụ.",
            "Sử dụng Ansible Vault để mã hóa các thông tin nhạy cảm như mật khẩu, API keys."
        ]
    },
    "Docker": {
        "task_title": "Đóng gói Ứng dụng với Docker Container",
        "sub_tasks": [
            "Viết Dockerfile tối ưu (sử dụng Multi-stage builds) để đóng gói mã nguồn thành Container Image siêu nhẹ.",
            "Sử dụng Docker Compose để cấu hình mạng kết nối nội bộ giữa web server, database và caching layers.",
            "Thực hành quản lý vòng đời container: logs, port mapping, data volumes để lưu trữ dữ liệu bền vững."
        ]
    },
    "Kubernetes": {
        "task_title": "Điều phối Container với Kubernetes (K8s)",
        "sub_tasks": [
            "Cài đặt Minikube cục bộ, làm quen với công cụ CLI kubectl để tương tác với cluster.",
            "Viết file YAML để khai báo các đối tượng Kubernetes cơ bản: Pods, Deployments, Services, ConfigMaps.",
            "Thiết lập cơ chế tự động mở rộng Pods (Horizontal Pod Autoscaler) và kiểm tra khả năng tự phục hồi khi pod bị sập."
        ]
    },
    "CI/CD": {
        "task_title": "Xây dựng Quy trình Tự động hóa CI/CD Pipeline",
        "sub_tasks": [
            "Tự động hóa quy trình kiểm thử (Unit test, Linting) sử dụng GitHub Actions, GitLab CI hoặc Jenkins.",
            "Cấu hình luồng Continuous Delivery tự động build docker image và đẩy lên Docker Hub / AWS ECR.",
            "Triển khai tự động mã nguồn mới lên các dịch vụ hosting đám mây (Render, AWS, Heroku) khi merge code vào nhánh main."
        ]
    },
    "Terraform": {
        "task_title": "Quản lý Hạ tầng bằng Mã nguồn (Terraform IaC)",
        "sub_tasks": [
            "Viết mã nguồn Terraform (.tf) để khai báo các tài nguyên hạ tầng đám mây (EC2, S3 bucket).",
            "Hiểu cách hoạt động của tệp trạng thái terraform.tfstate và cơ chế Locking khi làm việc nhóm.",
            "Sử dụng các lệnh cơ bản: terraform init, plan, apply, destroy để triển khai hạ tầng ảo hóa."
        ]
    },
    "Python": {
        "task_title": "Lập trình Python Nâng cao & Phân tích",
        "sub_tasks": [
            "Viết các kịch bản lệnh Python để tự động hóa sao lưu thư mục hoặc gọi API thu thập dữ liệu định kỳ.",
            "Sử dụng các thư viện phân tích dữ liệu chuyên nghiệp (pandas, numpy, matplotlib) để phân tích tập dữ liệu thực tế.",
            "Thực hiện kiểm thử mã nguồn bằng thư viện unittest hoặc pytest để đảm bảo tính đúng đắn của logic."
        ]
    },
    "Git": {
        "task_title": "Quản lý Phiên bản & Cộng tác với Git/GitHub",
        "sub_tasks": [
            "Khởi tạo repository Git cục bộ và đẩy lên GitHub, cấu hình tệp .gitignore bỏ qua thư mục rác.",
            "Thực hành quy trình làm việc Git Flow: tạo branch tính năng, commit rõ nghĩa và gửi Pull Request.",
            "Giải quyết các xung đột mã nguồn (Merge Conflicts) trực tiếp từ editor và hiểu cách quay lui lịch sử git reset/revert."
        ]
    },
    "Linux": {
        "task_title": "Quản trị Hệ điều hành Linux & Bash Shell",
        "sub_tasks": [
            "Làm quen với các lệnh Terminal cơ bản quản lý file và thư mục: ls, cd, mkdir, cp, mv, rm, find.",
            "Quản lý phân quyền truy cập tệp tin bằng lệnh chmod, chown và giám sát tiến trình hệ thống bằng top, htop, ps.",
            "Viết các kịch bản Bash Shell Script đơn giản chứa các câu lệnh điều kiện (if/else), vòng lặp (for/while)."
        ]
    }
}

def generate_local_roadmap(role, missing_skills):
    """Fallback function to divide required skills dynamically based on target role into 4 domains."""
    # Load all required skills for this role from job_to_skill.json
    try:
        with open("data/dataset/job_to_skill.json", "r", encoding="utf-8") as f:
            job_skills_map = json.load(f)
        role_skills = job_skills_map.get(role, [])
    except Exception:
        role_skills = []
        
    if not role_skills:
        role_skills = list(missing_skills) if missing_skills else ["Python", "Git", "SQL", "Docker", "CI/CD", "AWS"]

    # Let's group all role_skills into 4 thematic domains
    domain_1_skills = []
    domain_2_skills = []
    domain_3_skills = []
    domain_4_skills = []
    
    # Categorization keywords
    cat_1_keywords = ["git", "linux", "unix", "windows", "macos", "bash", "shell", "terminal", "command line", "networking", "protocol", "ssh", "ssl", "tls", "http", "dns", "security", "governance", "research", "methodology", "writing"]
    cat_2_keywords = ["python", "javascript", "typescript", "java", "c++", "c#", "go", "golang", "rust", "ruby", "php", "html", "css", "react", "angular", "vue", "node", "express", "django", "spring", "flask", "swift", "kotlin", "frontend", "backend", "web development", "spark", "hadoop", "mapreduce", "etl", "machine learning", "deep learning", "pytorch", "tensorflow", "algorithms", "data science"]
    cat_3_keywords = ["sql", "mysql", "postgresql", "mongodb", "database", "redis", "nosql", "api", "rest", "graphql", "system design", "cache", "caching", "architecture", "query", "queries", "warehouse", "lake", "modeling", "design"]
    cat_4_keywords = ["docker", "kubernetes", "k8s", "aws", "azure", "gcp", "google cloud", "cloud", "ci/cd", "github actions", "nginx", "apache", "serverless", "lambda", "deploy", "deployment", "devops", "ansible", "terraform", "jenkins", "puppet", "chef"]
    
    for s in role_skills:
        sl = s.lower()
        if any(k in sl for k in cat_1_keywords):
            domain_1_skills.append(s)
        elif any(k in sl for k in cat_2_keywords):
            domain_2_skills.append(s)
        elif any(k in sl for k in cat_3_keywords):
            domain_3_skills.append(s)
        elif any(k in sl for k in cat_4_keywords):
            domain_4_skills.append(s)
        else:
            # Leftovers round-robin
            total_added = len(domain_1_skills) + len(domain_2_skills) + len(domain_3_skills) + len(domain_4_skills)
            m = total_added % 4
            if m == 0:
                domain_1_skills.append(s)
            elif m == 1:
                domain_2_skills.append(s)
            elif m == 2:
                domain_3_skills.append(s)
            else:
                domain_4_skills.append(s)

    def build_tasks_for_skills(skills, fallback_domain_idx):
        tasks = []
        for s in skills:
            if s in SKILL_DETAILS_DB:
                tasks.append({
                    "task_title": SKILL_DETAILS_DB[s]["task_title"],
                    "sub_tasks": SKILL_DETAILS_DB[s]["sub_tasks"]
                })
            else:
                # Dynamic generic task generation for niche skills
                tasks.append({
                    "task_title": f"Tìm hiểu & Thực hành {s}",
                    "sub_tasks": [
                        f"Nghiên cứu tài liệu chính thức và nắm vững lý thuyết nền tảng của công nghệ {s}.",
                        f"Thiết lập môi trường phát triển cục bộ và viết một chương trình demo nhỏ áp dụng {s}.",
                        f"Tìm hiểu các tình huống thực tế thường sử dụng {s} trong doanh nghiệp."
                    ]
                })
        
        # If no tasks created for this domain, insert a general domain task
        if not tasks:
            default_tasks = [
                [
                    {"task_title": "Thiết lập Môi trường & Shell cơ bản", "sub_tasks": ["Cài đặt IDE và tiện ích mở rộng.", "Sử dụng dòng lệnh điều hướng thư mục.", "Quản lý mã nguồn với Git."]}
                ],
                [
                    {"task_title": "Lập trình Cốt lõi & Giải thuật", "sub_tasks": ["Nắm vững cấu trúc dữ liệu mảng, danh sách.", "Viết hàm xử lý logic và xử lý ngoại lệ.", "Luyện tập các bài toán logic thuật toán cơ bản."]}
                ],
                [
                    {"task_title": "Thiết kế & Lưu trữ Dữ liệu", "sub_tasks": ["Phân biệt cơ sở dữ liệu SQL và NoSQL.", "Thiết kế lược đồ quan hệ thực thể ERD.", "Viết các câu lệnh CRUD kết nối dữ liệu."]}
                ],
                [
                    {"task_title": "Container hóa & Tự động hóa CI/CD", "sub_tasks": ["Đóng gói phần mềm sử dụng Dockerfile.", "Viết kịch bản tự động hóa quy trình test.", "Triển khai ứng dụng lên máy chủ đám mây ảo."]}
                ]
            ]
            tasks = default_tasks[fallback_domain_idx]
        return tasks

    roadmap = [
        {
            "phase": "Vùng 1: Kiến thức Nền tảng & Hệ thống (Nền tảng)",
            "title": "Hệ điều hành, Nền tảng & Quy chuẩn bảo mật",
            "description": f"Thiết lập môi trường hệ điều hành và các giao thức bảo mật/quản lý cơ bản cho vai trò {role}.",
            "tasks": build_tasks_for_skills(domain_1_skills, 0)
        },
        {
            "phase": "Vùng 2: Phát triển & Logic xử lý (Ngôn ngữ & Xử lý)",
            "title": "Ngôn ngữ Lập trình & Logic tính toán",
            "description": f"Xây dựng tư duy giải thuật, cú pháp ngôn ngữ và các mô hình tính toán nghiệp vụ chính cho vai trò {role}.",
            "tasks": build_tasks_for_skills(domain_2_skills, 1)
        },
        {
            "phase": "Vùng 3: Cơ sở Dữ liệu & Thiết kế Hệ thống (Dữ liệu & API)",
            "title": "Cơ sở Dữ liệu, Kho lưu trữ & Thiết kế API",
            "description": f"Thiết kế cấu trúc bảng, kho lưu trữ dữ liệu và xây dựng các giao diện kết nối API cho vai trò {role}.",
            "tasks": build_tasks_for_skills(domain_3_skills, 2)
        },
        {
            "phase": "Vùng 4: Đám mây & Vận hành (Cloud & DevOps)",
            "title": "Container hóa, Triển khai Đám mây & Tự động hóa",
            "description": f"Đóng gói ảo hóa, tự động hóa quy trình CI/CD và vận hành hạ tầng đám mây cho vai trò {role}.",
            "tasks": build_tasks_for_skills(domain_4_skills, 3)
        }
    ]
    return roadmap

def generate_ai_roadmap(role, missing_skills, api_key):
    """Google Gemini-powered roadmap generator tailored to the candidate's missing skills."""
    if not api_key:
        return generate_local_roadmap(role, missing_skills)
        
    model_name = get_suitable_gemini_model(api_key)
    if not model_name:
        return generate_local_roadmap(role, missing_skills)
        
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(model_name)
        
        prompt = f"""
        You are an expert career coach and technical educator. Generate a highly detailed, thematic 4-domain learning roadmap for someone aiming to become a "{role}".
        The candidate is missing the following skills: {', '.join(missing_skills) if missing_skills else 'None (they want to master advanced skills)'}.
        IMPORTANT: The entire roadmap (phase titles, description, task titles, and sub-tasks) MUST be written in Vietnamese.
        
        Provide a highly specific, logical learning sequence structured as a **single valid JSON array** of objects, where each object represents a thematic knowledge domain (phân vùng kiến thức).
        
        Strict format:
        [
          {{
            "phase": "Vùng 1: Kiến thức Nền tảng & Hệ điều hành (Nền tảng)",
            "title": "Tên chuyên đề cụ thể liên quan đến hệ điều hành/môi trường (Ví dụ: Hệ điều hành, Shell & Quản lý Mã nguồn)...",
            "description": "Tóm tắt ngắn gọn những gì miền kiến thức này bao quát bằng tiếng Việt...",
            "tasks": [
              {{
                "task_title": "Tên chủ đề/kỹ năng cụ thể...",
                "sub_tasks": [
                  "Nhiệm vụ thực tế 1 bằng tiếng Việt...",
                  "Nhiệm vụ thực tế 2 bằng tiếng Việt..."
                ]
              }},
              ...
            ]
          }},
          ...
        ]
        
        Rules:
        - Return ONLY the JSON array. Do not include markdown code block syntax (like ```json), commentary, or introductions.
        - Ensure it is a valid parseable JSON.
        - Divide the roadmap into exactly 4 logical domains matching this sequence:
          1. "Vùng 1: Kiến thức Nền tảng & Hệ điều hành (Nền tảng)"
          2. "Vùng 2: Ngôn ngữ & Lập trình Cốt lõi (Ngôn ngữ)"
          3. "Vùng 3: Cơ sở Dữ liệu & Thiết kế Hệ thống (Dữ liệu & API)"
          4. "Vùng 4: DevOps, Điện toán Đám mây & Triển khai (DevOps & Cloud)"
        - In each domain, design detailed tasks and sub_tasks that are highly practical and tailored to master the missing skills listed above.
        - CRITICAL: Sub-tasks MUST be extremely specific and detailed. Avoid generic terms. Instead of "Học Python cơ bản", write specific instructions like "Viết mã nguồn Python sử dụng pandas để làm sạch dữ liệu từ tệp CSV và lưu trữ vào tệp JSON". Specify concrete tools, commands, frameworks, projects, and websites to read.
        - Ensure each domain contains at least 2 detailed tasks, and each task contains at least 3 detailed, step-by-step sub-tasks.
        """
        response = model.generate_content(prompt)
        if response and response.text:
            raw_text = response.text.strip()
            start = raw_text.find("[")
            end = raw_text.rfind("]") + 1
            if start != -1 and end != -1:
                raw_text = raw_text[start:end]
            return json.loads(raw_text)
    except Exception:
        return generate_local_roadmap(role, missing_skills)
    return generate_local_roadmap(role, missing_skills)
